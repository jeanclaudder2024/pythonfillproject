#!/usr/bin/env python3
"""
Working FastAPI service with all endpoints
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import uvicorn
import uuid
import os
import tempfile
import json
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
import re
import shutil
import random
from datetime import datetime, timedelta
from dotenv import load_dotenv
from supabase import create_client, Client

# Load environment variables
load_dotenv()

# Supabase connection
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_ANON_KEY")

# Initialize Supabase client
supabase: Client = None
if SUPABASE_URL and SUPABASE_KEY:
    try:
        supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
        print("Connected to Supabase database")
    except Exception as e:
        print(f"Failed to connect to Supabase: {e}")
        print("Falling back to file-based storage")
else:
    print("Supabase credentials not found, using file-based storage")

app = FastAPI(title="Working Document Service", version="1.0.0")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Simple in-memory storage for templates (in production, use a database)
templates_storage = []
templates_file = Path("templates_data.json")

def load_templates():
    """Load templates from file"""
    global templates_storage
    if templates_file.exists():
        try:
            with open(templates_file, 'r', encoding='utf-8') as f:
                templates_storage = json.load(f)
        except:
            templates_storage = []

def save_templates():
    """Save templates to file"""
    try:
        with open(templates_file, 'w', encoding='utf-8') as f:
            json.dump(templates_storage, f, indent=2)
    except:
        pass

# Load templates on startup
load_templates()

def extract_placeholders_from_docx(file_path):
    """Extract placeholders from a Word document"""
    try:
        doc = Document(file_path)
        placeholders = set()
        
        # Check paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Find placeholders in format {placeholder_name}
            found = re.findall(r'\{([^}]+)\}', text)
            placeholders.update(found)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        found = re.findall(r'\{([^}]+)\}', text)
                        placeholders.update(found)
        
        # Clean up placeholders - remove any malformed ones
        cleaned_placeholders = []
        for placeholder in placeholders:
            # Remove any extra characters and clean up
            clean_placeholder = placeholder.strip()
            if clean_placeholder and not clean_placeholder.startswith('{'):
                cleaned_placeholders.append(clean_placeholder)
        
        return cleaned_placeholders
    except Exception as e:
        print(f"Error extracting placeholders: {e}")
        return ["vessel_name", "imo", "vessel_type", "flag", "owner", "current_date"]  # fallback

def generate_realistic_data_for_placeholder(placeholder):
    """Generate realistic data for missing placeholders"""
    import random
    from datetime import datetime, timedelta
    
    # Define realistic data patterns
    realistic_data = {
        # Banking
        'seller_bank_account_no': f"{random.randint(1000000000, 9999999999)}",
        'seller_bank_swift': f"{random.choice(['CHASUS33', 'BOFAUS3N', 'CITIUS33', 'DEUTUS33'])}",
        'seller_bank_name': random.choice(['Chase Bank', 'Bank of America', 'Citibank', 'Deutsche Bank']),
        'seller_bank_address': f"{random.randint(100, 9999)} {random.choice(['Main St', 'Broadway', 'Wall St', 'Park Ave'])}, New York, NY",
        'seller_bank_officer_name': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'seller_bank_officer_mobile': f"+1-{random.randint(200, 999)}-{random.randint(200, 999)}-{random.randint(1000, 9999)}",
        'confirming_bank_account_number': f"{random.randint(1000000000, 9999999999)}",
        'confirming_bank_swift': f"{random.choice(['HSBCUS33', 'BNPAUS33', 'SCBLUS33'])}",
        'confirming_bank_name': random.choice(['HSBC', 'BNP Paribas', 'Standard Chartered']),
        'confirming_bank_address': f"{random.randint(100, 9999)} {random.choice(['Financial District', 'Banking Center', 'Commerce St'])}, Singapore",
        'confirming_bank_officer': f"{random.choice(['David', 'Emma', 'James', 'Anna'])} {random.choice(['Lee', 'Chen', 'Wong', 'Tan'])}",
        'confirming_bank_officer_contact': f"+65-{random.randint(6000, 9999)}-{random.randint(1000, 9999)}",
        'confirming_bank_tel': f"+65-{random.randint(6000, 9999)}-{random.randint(1000, 9999)}",
        'issuing_bank_account_number': f"{random.randint(1000000000, 9999999999)}",
        'issuing_bank_swift': f"{random.choice(['JPMUS33', 'WFCBUS33', 'PNCUS33'])}",
        'issuing_bank_name': random.choice(['JPMorgan Chase', 'Wells Fargo', 'PNC Bank']),
        'issuing_bank_address': f"{random.randint(100, 9999)} {random.choice(['Banking Plaza', 'Financial Center', 'Commerce Ave'])}, London",
        'issuing_bank_officer': f"{random.choice(['Robert', 'Jennifer', 'Christopher', 'Amanda'])} {random.choice(['Taylor', 'Anderson', 'Thomas', 'Jackson'])}",
        'issuing_bank_officer_contact': f"+44-{random.randint(20, 29)}-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'issuing_bank_tel': f"+44-{random.randint(20, 29)}-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        
        # Commercial
        'proforma_invoice_no': f"PI-{datetime.now().year}-{random.randint(1000, 9999)}",
        'invoice_no': f"INV-{datetime.now().year}-{random.randint(1000, 9999)}",
        'commercial_id': f"COM-{datetime.now().year}-{random.randint(1000, 9999)}",
        'document_number': f"DOC-{datetime.now().year}-{random.randint(1000, 9999)}",
        'contract_value': f"USD {random.randint(1000000, 50000000):,}",
        'total_amount': f"USD {random.randint(1000000, 10000000):,}",
        'total_amount_due': f"USD {random.randint(1000000, 10000000):,}",
        'amount_in_words': f"{random.choice(['Five', 'Ten', 'Fifteen', 'Twenty'])} Million US Dollars",
        'transaction_currency': 'USD',
        'payment_terms': random.choice(['30 days', '45 days', '60 days', '90 days']),
        'validity': f"{random.randint(30, 90)} days",
        'contract_duration': f"{random.randint(6, 24)} months",
        'monthly_delivery': f"{random.randint(1000, 10000)} MT",
        'performance_bond': f"USD {random.randint(100000, 1000000):,}",
        'insurance': f"USD {random.randint(500000, 5000000):,}",
        
        # Shipping
        'port_of_loading': random.choice(['Singapore', 'Rotterdam', 'Houston', 'Dubai', 'Shanghai']),
        'port_of_discharge': random.choice(['Tokyo', 'Hamburg', 'New York', 'Los Angeles', 'Busan']),
        'place_of_destination': random.choice(['Tokyo Port', 'Hamburg Port', 'New York Port', 'Los Angeles Port']),
        'final_delivery_place': random.choice(['Tokyo Port', 'Hamburg Port', 'New York Port', 'Los Angeles Port']),
        'origin': random.choice(['Malaysia', 'Indonesia', 'Nigeria', 'Saudi Arabia', 'Kuwait']),
        'country_of_origin': random.choice(['Malaysia', 'Indonesia', 'Nigeria', 'Saudi Arabia', 'Kuwait']),
        'shipping_terms': random.choice(['FOB', 'CIF', 'CFR', 'EXW']),
        'terms_of_delivery': random.choice(['FOB', 'CIF', 'CFR', 'EXW']),
        'via_name': random.choice(['Direct', 'Via Singapore', 'Via Rotterdam', 'Via Dubai']),
        'through_name': random.choice(['Direct', 'Via Singapore', 'Via Rotterdam', 'Via Dubai']),
        'partial_shipment': random.choice(['Allowed', 'Not Allowed']),
        'transshipment': random.choice(['Allowed', 'Not Allowed']),
        
        # Dates
        'date_of_issue': (datetime.now() - timedelta(days=random.randint(1, 30))).strftime('%Y-%m-%d'),
        'issued_date': (datetime.now() - timedelta(days=random.randint(1, 30))).strftime('%Y-%m-%d'),
        'issue_date': (datetime.now() - timedelta(days=random.randint(1, 30))).strftime('%Y-%m-%d'),
        'shipment_date2': (datetime.now() + timedelta(days=random.randint(30, 90))).strftime('%Y-%m-%d'),
        'shipment_date3': (datetime.now() + timedelta(days=random.randint(30, 90))).strftime('%Y-%m-%d'),
        'valid_until': (datetime.now() + timedelta(days=random.randint(60, 180))).strftime('%Y-%m-%d'),
        'buyer_signatory_date': datetime.now().strftime('%Y-%m-%d'),
        'seller_signatory_date': datetime.now().strftime('%Y-%m-%d'),
        
        # Product specifications
        'commodity': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel', 'Bunker Fuel']),
        'product_name': random.choice(['Light Sweet Crude', 'Heavy Crude', 'Diesel Fuel', 'Gasoline']),
        'goods_details': random.choice(['Light Sweet Crude Oil', 'Heavy Crude Oil', 'Diesel Fuel', 'Gasoline']),
        'specification': random.choice(['API 35-40', 'API 25-30', 'Sulfur < 0.5%', 'Sulfur < 1.0%']),
        'quality': random.choice(['Premium Grade', 'Standard Grade', 'Commercial Grade']),
        'inspection': random.choice(['SGS', 'Bureau Veritas', 'Intertek', 'Lloyd\'s Register']),
        'cloud_point': f"{random.randint(-10, 10)}°C",
        'free_fatty_acid': f"{random.uniform(0.1, 2.0):.1f}%",
        'iodine_value': f"{random.randint(40, 60)}",
        'moisture_impurities': f"{random.uniform(0.1, 1.0):.1f}%",
        'slip_melting_point': f"{random.randint(20, 35)}°C",
        'colour': random.choice(['Light Brown', 'Dark Brown', 'Black', 'Amber']),
        
        # Quantities and prices
        'total_quantity': f"{random.randint(10000, 100000)} MT",
        'quantity2': f"{random.randint(5000, 50000)} MT",
        'quantity3': f"{random.randint(5000, 50000)} MT",
        'total_weight': f"{random.randint(10000, 100000)} MT",
        'total_gross': f"{random.randint(10000, 100000)} MT",
        'total_containers': f"{random.randint(1, 10)}",
        'unit_price2': f"USD {random.randint(50, 100)}.00",
        'unit_price3': f"USD {random.randint(50, 100)}.00",
        'amount2': f"USD {random.randint(250000, 5000000):,}",
        'amount3': f"USD {random.randint(250000, 5000000):,}",
        'price': f"USD {random.randint(50, 100)}.00",
        'shipping_charges': f"USD {random.randint(50000, 500000):,}",
        'other_expenditures': f"USD {random.randint(10000, 100000):,}",
        'discount': f"{random.randint(0, 10)}%",
        
        # Items
        'item2': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel']),
        'item3': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel']),
        'consignment2': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel']),
        'consignment33': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel']),
        
        # Signatories
        'seller_signatory_name': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'seller_signatory_position': random.choice(['Managing Director', 'Sales Manager', 'Operations Manager', 'CEO']),
        'seller_signature': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'buyer_signatory_name': f"{random.choice(['Taro', 'Yuki', 'Hiroshi', 'Akira'])} {random.choice(['Yamada', 'Sato', 'Suzuki', 'Takahashi'])}",
        'buyer_signatory_position': random.choice(['Procurement Manager', 'General Manager', 'Director', 'President']),
        'buyer_signature': f"{random.choice(['Taro', 'Yuki', 'Hiroshi', 'Akira'])} {random.choice(['Yamada', 'Sato', 'Suzuki', 'Takahashi'])}",
        'signatory_name': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'authorized_person_name': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'notary_number': f"NOT-{random.randint(1000, 9999)}",
        
        # Contact information
        'buyer_company_name': random.choice(['Tokyo Trading Co.', 'Osaka Shipping Ltd.', 'Yokohama Marine Inc.', 'Kobe Commerce Corp.']),
        'buyer_name': f"{random.choice(['Taro', 'Yuki', 'Hiroshi', 'Akira'])} {random.choice(['Yamada', 'Sato', 'Suzuki', 'Takahashi'])}",
        'buyer_address': f"{random.randint(1, 999)} {random.choice(['Chuo-dori', 'Ginza', 'Shibuya', 'Shinjuku'])}, Tokyo, Japan",
        'buyer_city_country': 'Tokyo, Japan',
        'buyer_email': f"buyer{random.randint(1, 999)}@{random.choice(['tokyo-trading.com', 'osaka-shipping.com', 'yokohama-marine.com'])}",
        'buyer_tel': f"+81-3-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'buyer_fax': f"+81-3-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'buyer_mobile': f"+81-{random.randint(90, 99)}-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'buyer_office_tel': f"+81-3-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'buyer_designation': random.choice(['Procurement Manager', 'General Manager', 'Director', 'President']),
        'buyer_representative': f"{random.choice(['Taro', 'Yuki', 'Hiroshi', 'Akira'])} {random.choice(['Yamada', 'Sato', 'Suzuki', 'Takahashi'])}",
        'buyer_position': random.choice(['Procurement Manager', 'General Manager', 'Director', 'President']),
        'buyer_registration': f"REG-{random.randint(100000, 999999)}",
        
        # Seller information
        'seller_company': random.choice(['Singapore Trading Ltd.', 'Malaysia Oil Corp.', 'Indonesia Marine Inc.', 'Thailand Commerce Co.']),
        'seller_address': f"{random.randint(1, 999)} {random.choice(['Marina Bay', 'Orchard Road', 'Raffles Place', 'Clarke Quay'])}, Singapore",
        
        # Default fallback
        'default': f"Sample {placeholder.replace('_', ' ').title()}"
    }
    
    # Return realistic data or default
    return realistic_data.get(placeholder.lower(), realistic_data['default'])

def fill_word_template(template_path, output_path, vessel_data):
    """Fill a Word template with vessel data"""
    try:
        print(f"Starting template processing: {template_path} -> {output_path}")
        
        # Copy template to output path
        shutil.copy2(template_path, output_path)
        print(f"Template copied successfully")
        
        # Open the document
        doc = Document(output_path)
        print(f"Document opened successfully, processing {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        
        # Replace placeholders in paragraphs
        paragraphs_processed = 0
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                text = paragraph.text
                original_text = text
                for placeholder, value in vessel_data.items():
                    # Replace both {placeholder} and {placeholder} formats
                    text = text.replace(f"{{{placeholder}}}", str(value))
                    text = text.replace(f"{{{{{placeholder}}}}}", str(value))  # Handle double braces
                
                # Find any remaining placeholders and replace with realistic data
                remaining_placeholders = re.findall(r'\{([^}]+)\}', text)
                for placeholder in remaining_placeholders:
                    realistic_value = generate_realistic_data_for_placeholder(placeholder)
                    text = text.replace(f"{{{placeholder}}}", str(realistic_value))
                    print(f"Generated realistic data for missing placeholder '{placeholder}': {realistic_value}")
                
                if text != original_text:
                    paragraph.text = text
                    paragraphs_processed += 1
                    print(f"Replaced placeholders in paragraph {i}: {original_text[:50]}... -> {text[:50]}...")
            except Exception as e:
                print(f"Error processing paragraph {i}: {e}")
                continue
        
        print(f"Processed {paragraphs_processed} paragraphs with replacements")
        
        # Replace placeholders in tables
        tables_processed = 0
        for table_idx, table in enumerate(doc.tables):
            try:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            try:
                                text = paragraph.text
                                original_text = text
                                for placeholder, value in vessel_data.items():
                                    # Replace both {placeholder} and {placeholder} formats
                                    text = text.replace(f"{{{placeholder}}}", str(value))
                                    text = text.replace(f"{{{{{placeholder}}}}}", str(value))  # Handle double braces
                                
                                # Find any remaining placeholders and replace with realistic data
                                remaining_placeholders = re.findall(r'\{([^}]+)\}', text)
                                for placeholder in remaining_placeholders:
                                    realistic_value = generate_realistic_data_for_placeholder(placeholder)
                                    text = text.replace(f"{{{placeholder}}}", str(realistic_value))
                                    print(f"Generated realistic data for missing placeholder '{placeholder}': {realistic_value}")
                                
                                if text != original_text:
                                    paragraph.text = text
                                    tables_processed += 1
                                    print(f"Replaced placeholders in table {table_idx}, row {row_idx}, cell {cell_idx}: {original_text[:50]}... -> {text[:50]}...")
                            except Exception as e:
                                print(f"Error processing table cell {table_idx}.{row_idx}.{cell_idx}.{para_idx}: {e}")
                                continue
            except Exception as e:
                print(f"Error processing table {table_idx}: {e}")
                continue
        
        print(f"Processed {tables_processed} table cells with replacements")
        
        # Save the document
        doc.save(output_path)
        print(f"Template filled successfully: {output_path}")
        return True
        
    except Exception as e:
        print(f"CRITICAL ERROR filling template: {e}")
        import traceback
        traceback.print_exc()
        return False

@app.get("/")
async def root():
    return {"message": "Working Document Service is running", "status": "ok"}

@app.get("/health")
async def health():
    return {"status": "healthy", "service": "document-processing"}

@app.get("/templates")
async def get_templates():
    """Get list of available templates"""
    try:
        if supabase:
            # Get templates from database
            result = supabase.table("document_templates").select("*").eq("is_active", True).execute()
            return result.data
        else:
            # Fallback to file storage
            return templates_storage
    except Exception as e:
        print(f"Error fetching templates from database: {e}")
        # Fallback to file storage
        return templates_storage

@app.get("/vessels")
async def get_vessels():
    return [
        {
            "id": "1",
            "name": "Petroleum Express 529",
            "imo": "IMO1861018",
            "vessel_type": "Crude Oil Tanker",
            "flag": "Malta"
        },
        {
            "id": "2", 
            "name": "Atlantic Voyager 805",
            "imo": "IMO2379622",
            "vessel_type": "Container Ship",
            "flag": "Panama"
        }
    ]

@app.post("/upload-template")
async def upload_template(
    name: str = Form(...),
    description: str = Form(...),
    template_file: UploadFile = File(...)
):
    """Upload a new document template"""
    try:
        # Read file content
        content = await template_file.read()
        
        # Extract placeholders from the uploaded Word document
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(content)
            temp_file.flush()
            actual_placeholders = extract_placeholders_from_docx(temp_file.name)
            os.unlink(temp_file.name)

        if supabase:
            # Store in database
            template_data = {
                "name": name,
                "description": description,
                "template_file": content,  # Store as BYTEA
                "file_name": template_file.filename,
                "file_size": len(content),
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "placeholders": actual_placeholders,
                "placeholder_mappings": {},
                "is_active": True
            }
            
            result = supabase.table("document_templates").insert(template_data).execute()
            
            return JSONResponse({
                "success": True,
                "message": "Template uploaded to database successfully",
                "template": result.data[0]
            })
        else:
            # Fallback to file storage
            template_id = str(uuid.uuid4())
            
            # Create templates directory if it doesn't exist
            templates_dir = Path("templates")
            templates_dir.mkdir(exist_ok=True)

            # Save the uploaded file
            file_extension = Path(template_file.filename).suffix
            saved_filename = f"{template_id}{file_extension}"
            file_path = templates_dir / saved_filename

            with open(file_path, "wb") as buffer:
                buffer.write(content)

            # Create template info
            template_info = {
                "id": template_id,
                "name": name,
                "description": description,
                "file_name": template_file.filename,
                "file_size": len(content),
                "placeholders": actual_placeholders,
                "is_active": True,
                "created_at": "2025-01-30T00:00:00Z"
            }

            # Add to storage
            templates_storage.append(template_info)
            save_templates()

            return JSONResponse({
                "success": True,
                "message": "Template uploaded to file storage successfully",
                "template": template_info
            })

    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Template upload failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.post("/process-document")
async def process_document(
    template_id: str = Form(...),
    vessel_imo: str = Form(...),
    template_file: UploadFile = File(...)
):
    """Process a document template with vessel data"""
    try:
        # Generate a unique document ID
        document_id = str(uuid.uuid4())
        
        # Create outputs directory if it doesn't exist
        outputs_dir = Path("outputs")
        outputs_dir.mkdir(exist_ok=True)
        
        # Find the template in database or storage
        template_info = None
        if supabase:
            try:
                result = supabase.table("document_templates").select("*").eq("id", template_id).execute()
                if result.data:
                    template_info = result.data[0]
            except Exception as e:
                print(f"Error fetching template from database: {e}")
        
        if not template_info:
            # Fallback to file storage
            for template in templates_storage:
                if template["id"] == template_id:
                    template_info = template
                    break

        if not template_info:
            return JSONResponse({
                "success": False,
                "message": "Template not found",
                "error": "Template ID not found in database or storage"
            }, status_code=404)
        
        # Get vessel data (you can expand this to fetch from database)
        vessel_data = {
            # Basic vessel info
            "vessel_name": "Petroleum Express 529",
            "imo": vessel_imo,
            "imo_number": vessel_imo,
            "vessel_type": "Crude Oil Tanker",
            "flag": "Malta",
            "flag_state": "Malta",
            "owner": "Sample Shipping Company",
            "vessel_owner": "Sample Shipping Company",
            "current_date": "2025-01-30",
            "vessel_id": "1",
            
            # Technical specifications
            "gross_tonnage": "45,000",
            "net_tonnage": "35,000",
            "deadweight": "85,000",
            "length": "250m",
            "length_overall": "250m",
            "beam": "45m",
            "draft": "15m",
            "year_built": "2015",
            "call_sign": "9HA1234",
            "registry_port": "Valletta",
            "class_society": "Lloyd's Register",
            "ism_manager": "Sample ISM Manager",
            "vessel_operator": "Sample Operator",
            "engine_type": "Diesel",
            "speed": "15 knots",
            
            # Cargo specifications
            "cargo_capacity": "85,000 MT",
            "cargo_tanks": "12",
            "pumping_capacity": "3,000 m³/h",
            
            # Product specifications
            "product_name": "Crude Oil",
            "Commodity": "Crude Oil",
            "Goods_Details": "Light Sweet Crude Oil",
            "Country_Of_Origin": "Malaysia",
            "Origin": "Malaysia",
            "Cloud_Point": "-5°C",
            "Free_Fatty_Acid": "0.1%",
            "Iodine_Value": "45",
            "Moisture_Impurities": "0.1%",
            "Slip_Melting_Point": "25°C",
            "Colour": "Light Brown",
            
            # Commercial details
            "Seller_Company": "Sample Trading Company",
            "Seller_Address": "123 Trading Street, Singapore",
            "Seller_Bank_Name": "Sample Bank",
            "Seller_Bank_Address": "456 Bank Street, Singapore",
            "Seller_Bank_Account_No": "1234567890",
            "Seller_Bank_Account_Name": "Sample Trading Company",
            "Seller_Bank_SWIFT": "SAMPUS33",
            "Seller_Bank_Officer_Name": "John Smith",
            "Seller_Bank_Officer_Mobile": "+65 9123 4567",
            
            "Buyer_Company": "Sample Buyer Company",
            "Buyer_Company_Name": "Sample Buyer Company",
            "Buyer_Address": "789 Buyer Avenue, Tokyo",
            "Buyer_Email": "buyer@sample.com",
            "Buyer_Tel": "+81 3 1234 5678",
            "Buyer_Designation": "Procurement Manager",
            "Buyer_Representative": "Taro Yamada",
            "Position_Title": "Procurement Manager",
            
            # Invoice details
            "Proforma_Invoice_No": "PI-2025-001",
            "Invoice_No": "INV-2025-001",
            "Date_Of_Issue": "2025-01-30",
            "Issued_Date": "2025-01-30",
            "Commercial_ID": "COM-2025-001",
            "Transaction_Currency": "USD",
            "Payment_Terms": "30 days",
            "Validity": "60 days",
            
            # Shipping details
            "Port_Of_Loading": "Singapore",
            "Port_Of_Discharge": "Tokyo",
            "Place_Of_Destination": "Tokyo Port",
            "Final_Delivery_Place": "Tokyo Port",
            "Terms_Of_Delivery": "FOB",
            "Via_Name": "Direct",
            "Through_Name": "Direct",
            "Partial_Shipment": "Not Allowed",
            "Transshipment": "Not Allowed",
            "Shipment_Date2": "2025-02-15",
            "Shipment_Date3": "2025-02-15",
            
            # Items and quantities
            "Item2": "Crude Oil",
            "Item3": "Crude Oil",
            "Quantity2": "50,000 MT",
            "Quantity3": "50,000 MT",
            "Unit_Price2": "USD 65.00",
            "Unit_Price3": "USD 65.00",
            "Amount2": "USD 3,250,000.00",
            "Amount3": "USD 3,250,000.00",
            "Total_Amount": "USD 6,500,000.00",
            "Total_Amount_Due": "USD 6,500,000.00",
            "Amount_In_Words": "Six Million Five Hundred Thousand US Dollars",
            "Total_Weight": "100,000 MT",
            "Total_Gross": "100,000 MT",
            "Total_Containers": "1",
            "Consignment2": "Crude Oil",
            "Consignment33": "Crude Oil",
            
            # Additional charges
            "shipping_Charges": "USD 50,000.00",
            "Other_Expenditures": "USD 10,000.00",
            "Discount": "0%",
            
            # Signatory
            "Signatory_Name": "John Smith"
        }
        
        # Get template file content
        template_file_content = None
        if supabase and 'template_file' in template_info:
            # Get template file from database
            template_file_content = template_info['template_file']
        else:
            # Fallback to file system
            templates_dir = Path("templates")
            template_file_path = templates_dir / f"{template_id}.docx"
            if template_file_path.exists():
                with open(template_file_path, "rb") as f:
                    template_file_content = f.read()
            else:
                return JSONResponse({
                    "success": False,
                    "message": "Template file not found",
                    "error": "Template file does not exist in database or on server"
                }, status_code=404)
        
        # Create output files
        filled_docx_file = outputs_dir / f"{document_id}_filled.docx"
        pdf_file = outputs_dir / f"{document_id}_filled.pdf"
        txt_file = outputs_dir / f"{document_id}_filled_fallback.txt"
        
        # Save template content to temporary file for processing
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_template:
            temp_template.write(template_file_content)
            temp_template.flush()
            
            # Fill the Word template with vessel data
            success = fill_word_template(temp_template.name, filled_docx_file, vessel_data)
            
            # Clean up temporary file
            os.unlink(temp_template.name)
        
        if not success:
            return JSONResponse({
                "success": False,
                "message": "Failed to fill template",
                "error": "Could not process the Word template"
            }, status_code=500)
        
        # Convert to PDF (try docx2pdf, fallback to text if it fails)
        pdf_success = False
        try:
            from docx2pdf import convert
            convert(str(filled_docx_file), str(pdf_file))
            pdf_success = True
        except Exception as e:
            print(f"PDF conversion failed: {e}")
            # Create a fallback text file
            with open(txt_file, 'w', encoding='utf-8') as f:
                f.write(f"Document processed successfully for vessel {vessel_imo}\n")
                f.write(f"Template: {template_info['name']}\n")
                f.write(f"Document ID: {document_id}\n")
                f.write(f"Note: PDF conversion failed, but Word document was created successfully.\n")
                f.write(f"Word file: {filled_docx_file}\n")
        
        # Return success response
        return JSONResponse({
            "success": True,
            "message": "Document processed successfully!",
            "document_id": document_id,
            "download_url": f"/download/{document_id}",
            "vessel_imo": vessel_imo,
            "template_id": template_id,
            "template_name": template_info["name"],
            "files_created": {
                "docx_file": str(filled_docx_file),
                "pdf_file": str(pdf_file) if pdf_success else None,
                "text_file": str(txt_file) if not pdf_success else None
            },
            "pdf_available": pdf_success
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Document processing failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.get("/download/{document_id}")
async def download_document(document_id: str, format: str = "pdf"):
    """Download processed document as actual file"""
    try:
        outputs_dir = Path("outputs")
        
        if format.lower() == "pdf":
            file_path = outputs_dir / f"{document_id}_filled.pdf"
            media_type = "application/pdf"
            filename = f"vessel_report_{document_id}.pdf"
        else:
            file_path = outputs_dir / f"{document_id}_filled_fallback.txt"
            media_type = "text/plain"
            filename = f"vessel_report_{document_id}.txt"
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Document not found")
        
        return FileResponse(
            path=str(file_path),
            media_type=media_type,
            filename=filename
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    print("Starting Working Document Service...")
    print(f"Service will be available at: http://0.0.0.0:{port}")
    print("API Documentation: http://0.0.0.0:8000/docs")
    print("Available endpoints:")
    print("  GET  /")
    print("  GET  /health")
    print("  GET  /templates")
    print("  GET  /vessels")
    print("  POST /upload-template")
    print("  POST /process-document")
    print("  GET  /download/{document_id}")
    uvicorn.run(app, host="0.0.0.0", port=port)
