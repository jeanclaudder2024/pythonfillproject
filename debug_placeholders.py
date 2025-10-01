#!/usr/bin/env python3
"""
Debug Placeholders Script
Examines the actual content of uploaded documents to see what placeholders they contain
"""

import os
from docx import Document
import re

def examine_document_content(file_path):
    """Examine the content of a document to see what placeholders it contains"""
    try:
        doc = Document(file_path)
        
        print(f"📄 Examining: {os.path.basename(file_path)}")
        print("=" * 60)
        
        # Extract all text from the document
        total_text = ""
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            total_text += paragraph.text + "\n"
        
        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_text += paragraph.text + "\n"
        
        print(f"📊 Total characters: {len(total_text)}")
        print()
        
        # Find different types of placeholders
        curly_single = re.findall(r'\{([^}]+)\}', total_text)
        curly_double = re.findall(r'\{\{([^}]+)\}\}', total_text)
        square_single = re.findall(r'\[([^\]]+)\]', total_text)
        square_double = re.findall(r'\[\[([^\]]+)\]\]', total_text)
        
        print("🔍 PLACEHOLDER ANALYSIS:")
        print(f"  Single curly braces {{placeholder}}: {len(curly_single)}")
        if curly_single:
            print(f"    Examples: {curly_single[:5]}")
        
        print(f"  Double curly braces {{{{placeholder}}}}: {len(curly_double)}")
        if curly_double:
            print(f"    Examples: {curly_double[:5]}")
        
        print(f"  Single square brackets [placeholder]: {len(square_single)}")
        if square_single:
            print(f"    Examples: {square_single[:5]}")
        
        print(f"  Double square brackets [[placeholder]]: {len(square_double)}")
        if square_double:
            print(f"    Examples: {square_double[:5]}")
        
        print()
        
        # Look for malformed patterns
        malformed_curly = re.findall(r'\{[^}]*$|^[^{]*\}', total_text, re.MULTILINE)
        malformed_square = re.findall(r'\[[^\]]*$|^[^\[]*\]', total_text, re.MULTILINE)
        incomplete_curly = re.findall(r'\{[A-Za-z0-9]*(?=\s|\n|$)', total_text)
        incomplete_square = re.findall(r'\[[A-Za-z0-9]*(?=\s|\n|$)', total_text)
        
        print("⚠️  MALFORMED PATTERNS:")
        print(f"  Malformed curly braces: {len(malformed_curly)}")
        if malformed_curly:
            print(f"    Examples: {malformed_curly[:3]}")
        
        print(f"  Malformed square brackets: {len(malformed_square)}")
        if malformed_square:
            print(f"    Examples: {malformed_square[:3]}")
        
        print(f"  Incomplete curly braces: {len(incomplete_curly)}")
        if incomplete_curly:
            print(f"    Examples: {incomplete_curly[:3]}")
        
        print(f"  Incomplete square brackets: {len(incomplete_square)}")
        if incomplete_square:
            print(f"    Examples: {incomplete_square[:3]}")
        
        print()
        
        # Show some sample text to understand the context
        print("📝 SAMPLE TEXT (first 500 characters):")
        print("-" * 40)
        print(total_text[:500])
        print("-" * 40)
        
        return {
            'curly_single': curly_single,
            'curly_double': curly_double,
            'square_single': square_single,
            'square_double': square_double,
            'total_text': total_text
        }
        
    except Exception as e:
        print(f"❌ Error examining document: {str(e)}")
        return None

def main():
    """Main function to examine documents"""
    print("🔍 DOCUMENT PLACEHOLDER DEBUGGER")
    print("=" * 60)
    
    # Look for a sample document to examine
    uploads_dir = "uploads"
    
    # Get the first few documents
    import glob
    docs = glob.glob(os.path.join(uploads_dir, "*.docx"))
    
    if not docs:
        print("❌ No documents found in uploads folder!")
        return
    
    print(f"📋 Found {len(docs)} documents. Examining first 3...")
    print()
    
    for i, doc_path in enumerate(docs[:3]):
        examine_document_content(doc_path)
        if i < 2:  # Add separator between documents
            print("\n" + "="*80 + "\n")

if __name__ == "__main__":
    main()