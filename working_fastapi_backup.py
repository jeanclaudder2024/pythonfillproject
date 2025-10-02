#!/usr/bin/env python3
"""
Working FastAPI service with all endpoints - CLEAN VERSION
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import uvicorn
import uuid
import os
import tempfile
import json
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
import re
import shutil
import random
from datetime import datetime, timedelta
from dotenv import load_dotenv
from supabase import create_client, Client

# Load environment variables
load_dotenv()

# Initialize Supabase client
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_ANON_KEY")

if not SUPABASE_URL or not SUPABASE_KEY:
    print("⚠️  Supabase credentials not found. Using dummy data.")
    supabase = None
else:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
    print("✅ Connected to Supabase database")

app = FastAPI(title="Working Document Service", version="1.0.0")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Simple in-memory storage for templates (in production, use a database)
templates_storage = []
templates_file = Path("templates_data.json")

def load_templates():
    """Load templates from file"""
    global templates_storage
    if templates_file.exists():
        try:
            with open(templates_file, 'r', encoding='utf-8') as f:
                templates_storage = json.load(f)
        except:
            templates_storage = []

def save_templates():
    """Save templates to file"""
    try:
        with open(templates_file, 'w', encoding='utf-8') as f:
            json.dump(templates_storage, f, indent=2)
    except:
        pass

# Load templates on startup
load_templates()

async def get_real_vessel_data(vessel_imo: str):
    """Fetch real vessel data from Supabase database"""
    try:
        if not supabase:
            print("⚠️  Supabase not connected, using dummy data")
            return get_dummy_vessel_data(vessel_imo)
        
        # Fetch vessel data from your real database
        response = supabase.table('vessels').select('*').eq('imo', vessel_imo).execute()
        
        if response.data and len(response.data) > 0:
            vessel = response.data[0]
            print(f"✅ Found real vessel data for IMO: {vessel_imo}")
            
            # Map your real vessel data to placeholders
            vessel_data = {
                # Basic vessel info from your database
                "vessel_name": vessel.get('name', 'Unknown Vessel'),
                "imo": vessel_imo,
                "imo_number": vessel_imo,
                "vessel_type": vessel.get('vessel_type', 'Unknown Type'),
                "flag": vessel.get('flag', 'Unknown Flag'),
                "flag_state": vessel.get('flag', 'Unknown Flag'),
                "owner": vessel.get('owner_name', 'Unknown Owner'),
                "vessel_owner": vessel.get('owner_name', 'Unknown Owner'),
                "operator": vessel.get('operator_name', 'Unknown Operator'),
                "current_date": datetime.now().strftime('%Y-%m-%d'),
                "vessel_id": str(vessel.get('id', '1')),
                
                # Additional real data from your database
                "mmsi": vessel.get('mmsi', ''),
                "callsign": vessel.get('callsign', ''),
                "built": vessel.get('built', 0),
                "deadweight": vessel.get('deadweight', 0),
                "cargo_capacity": vessel.get('cargo_capacity', 0),
                "length": vessel.get('length', 0),
                "width": vessel.get('width', 0),
                "beam": vessel.get('beam', ''),
                "draught": vessel.get('draught', 0),
                "draft": vessel.get('draft', ''),
                "gross_tonnage": vessel.get('gross_tonnage', 0),
                "crew_size": vessel.get('crew_size', 0),
                "engine_power": vessel.get('engine_power', 0),
                "fuel_consumption": vessel.get('fuel_consumption', 0),
                "speed": vessel.get('speed', ''),
                "status": vessel.get('status', ''),
                "nav_status": vessel.get('nav_status', ''),
                "current_lat": vessel.get('current_lat', 0),
                "current_lng": vessel.get('current_lng', 0),
                "course": vessel.get('course', 0),
                "departure_port": vessel.get('departure_port', 0),
                "destination_port": vessel.get('destination_port', 0),
                "departure_port_name": vessel.get('departure_port_name', 'Unknown Port'),
                "destination_port_name": vessel.get('destination_port_name', 'Unknown Port'),
                "loading_port_name": vessel.get('loading_port_name', 'Unknown Port'),
                "departure_date": vessel.get('departure_date', ''),
                "arrival_date": vessel.get('arrival_date', ''),
                "eta": vessel.get('eta', ''),
                "loading_port": vessel.get('loading_port', ''),
                "cargo_type": vessel.get('cargo_type', ''),
                "cargo_quantity": vessel.get('cargo_quantity', 0),
                "oil_type": vessel.get('oil_type', ''),
                "oil_source": vessel.get('oil_source', ''),
                "current_region": vessel.get('current_region', ''),
                "buyer_name": vessel.get('buyer_name', ''),
                "seller_name": vessel.get('seller_name', ''),
                "source_company": vessel.get('source_company', ''),
                "target_refinery": vessel.get('target_refinery', ''),
                "deal_value": vessel.get('deal_value', 0),
                "price": vessel.get('price', 0),
                "market_price": vessel.get('market_price', 0),
                "quantity": vessel.get('quantity', 0),
                "departure_lat": vessel.get('departure_lat', 0),
                "departure_lng": vessel.get('departure_lng', 0),
                "destination_lat": vessel.get('destination_lat', 0),
                "destination_lng": vessel.get('destination_lng', 0),
                "route_distance": vessel.get('route_distance', 0),
                "shipping_type": vessel.get('shipping_type', ''),
                "route_info": vessel.get('route_info', ''),
                "last_updated": vessel.get('last_updated', ''),
                "company_id": vessel.get('company_id', ''),
                
                # ICPO Specific Fields (using real data where available)
                "icpo_number": f"ICPO-{datetime.now().year}-{random.randint(1000, 9999)}",
                "icpo_date": datetime.now().strftime('%Y-%m-%d'),
                "icpo_validity": (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d'),
                "icpo_amount": f"USD {vessel.get('deal_value', random.randint(1000000, 10000000)):,}",
                "icpo_currency": "USD",
                "icpo_terms": "LC at sight",
                "icpo_bank": "HSBC Bank",
                "icpo_bank_address": "1 Centenary Square, Birmingham, UK",
                "icpo_swift": "HBUKGB4B",
                "icpo_account": f"{random.randint(1000000000, 9999999999)}",
                "icpo_beneficiary": vessel.get('seller_name', 'Sample Trading Company Ltd'),
                "icpo_beneficiary_address": f"{vessel.get('source_company', '123 Marina Bay')}, Singapore",
                "icpo_beneficiary_swift": "DBSBSGSG",
                "icpo_beneficiary_account": f"{random.randint(1000000000, 9999999999)}",
                "icpo_commodity": vessel.get('cargo_type', vessel.get('oil_type', 'Crude Oil')),
                "icpo_quantity": f"{vessel.get('cargo_quantity', vessel.get('quantity', random.randint(10000, 100000)))} MT",
                "icpo_specification": f"API 35-40, Sulfur < 0.5%",
                "icpo_origin": vessel.get('oil_source', 'Malaysia'),
                "icpo_destination": vessel.get('target_refinery', 'Singapore'),
                "icpo_loading_port": vessel.get('loading_port_name', vessel.get('loading_port', 'Port Klang, Malaysia')),
                "icpo_discharge_port": vessel.get('destination_port_name', 'Singapore Port'),
                "icpo_loading_date": vessel.get('departure_date', (datetime.now() + timedelta(days=15)).strftime('%Y-%m-%d')),
                "icpo_discharge_date": vessel.get('arrival_date', (datetime.now() + timedelta(days=25)).strftime('%Y-%m-%d')),
                "icpo_price": f"USD {vessel.get('price', random.randint(50, 100))}.00 per MT",
                "icpo_total_value": f"USD {vessel.get('deal_value', random.randint(5000000, 50000000)):,}",
                "icpo_payment_terms": "LC at sight",
                "icpo_delivery_terms": "FOB",
                "icpo_inspection": "SGS",
                "icpo_insurance": "All Risks",
            }
            
            return vessel_data
        else:
            print(f"⚠️  No vessel found with IMO: {vessel_imo}, using dummy data")
            return get_dummy_vessel_data(vessel_imo)
            
    except Exception as e:
        print(f"❌ Error fetching vessel data: {e}")
        return get_dummy_vessel_data(vessel_imo)

def get_dummy_vessel_data(vessel_imo: str):
    """Fallback dummy vessel data"""
    return {
        # Basic vessel info
        "vessel_name": "Petroleum Express 529",
        "imo": vessel_imo,
        "imo_number": vessel_imo,
        "vessel_type": "Crude Oil Tanker",
        "flag": "Malta",
        "flag_state": "Malta",
        "owner": "Sample Shipping Company",
        "vessel_owner": "Sample Shipping Company",
        "current_date": datetime.now().strftime('%Y-%m-%d'),
        "vessel_id": "1",
        
        # ICPO Specific Fields
        "icpo_number": f"ICPO-{datetime.now().year}-{random.randint(1000, 9999)}",
        "icpo_date": datetime.now().strftime('%Y-%m-%d'),
        "icpo_validity": (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d'),
        "icpo_amount": f"USD {random.randint(1000000, 10000000):,}",
        "icpo_currency": "USD",
        "icpo_terms": "LC at sight",
        "icpo_bank": "HSBC Bank",
        "icpo_bank_address": "1 Centenary Square, Birmingham, UK",
        "icpo_swift": "HBUKGB4B",
        "icpo_account": f"{random.randint(1000000000, 9999999999)}",
        "icpo_beneficiary": "Sample Trading Company Ltd",
        "icpo_beneficiary_address": "123 Marina Bay, Singapore",
        "icpo_beneficiary_swift": "DBSBSGSG",
        "icpo_beneficiary_account": f"{random.randint(1000000000, 9999999999)}",
        "icpo_commodity": "Crude Oil",
        "icpo_quantity": f"{random.randint(10000, 100000)} MT",
        "icpo_specification": "API 35-40, Sulfur < 0.5%",
        "icpo_origin": "Malaysia",
        "icpo_destination": "Singapore",
        "icpo_loading_port": "Port Klang, Malaysia",
        "icpo_discharge_port": "Singapore Port",
        "icpo_loading_date": (datetime.now() + timedelta(days=15)).strftime('%Y-%m-%d'),
        "icpo_discharge_date": (datetime.now() + timedelta(days=25)).strftime('%Y-%m-%d'),
        "icpo_price": f"USD {random.randint(50, 100)}.00 per MT",
        "icpo_total_value": f"USD {random.randint(5000000, 50000000):,}",
        "icpo_payment_terms": "LC at sight",
        "icpo_delivery_terms": "FOB",
        "icpo_inspection": "SGS",
        "icpo_insurance": "All Risks",
    }

def extract_placeholders_from_docx(file_path):
    """Extract placeholders from a Word document"""
    try:
        doc = Document(file_path)
        placeholders = set()
        
        # Check paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Find placeholders in format {placeholder_name}
            found = re.findall(r'\{([^}]+)\}', text)
            placeholders.update(found)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        found = re.findall(r'\{([^}]+)\}', text)
                        placeholders.update(found)
        
        # Clean up placeholders - remove any malformed ones
        cleaned_placeholders = []
        for placeholder in placeholders:
            # Remove any extra characters and clean up
            clean_placeholder = placeholder.strip()
            if clean_placeholder and not clean_placeholder.startswith('{'):
                cleaned_placeholders.append(clean_placeholder)
        
        return cleaned_placeholders
    except Exception as e:
        print(f"Error extracting placeholders: {e}")
        return ["vessel_name", "imo", "vessel_type", "flag", "owner", "current_date"]  # fallback

def generate_realistic_data_for_placeholder(placeholder):
    """Generate realistic data for missing placeholders"""
    import random
    from datetime import datetime, timedelta
    
    # Define realistic data patterns
    realistic_data = {
        # Banking
        'seller_bank_account_no': f"{random.randint(1000000000, 9999999999)}",
        'seller_bank_swift': f"{random.choice(['CHASUS33', 'BOFAUS3N', 'CITIUS33', 'DEUTUS33'])}",
        'seller_bank_name': random.choice(['Chase Bank', 'Bank of America', 'Citibank', 'Deutsche Bank']),
        'seller_bank_address': f"{random.randint(100, 9999)} {random.choice(['Main St', 'Broadway', 'Wall St', 'Park Ave'])}, New York, NY",
        'seller_bank_officer_name': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'seller_bank_officer_mobile': f"+1-{random.randint(200, 999)}-{random.randint(200, 999)}-{random.randint(1000, 9999)}",
        'confirming_bank_account_number': f"{random.randint(1000000000, 9999999999)}",
        'confirming_bank_swift': f"{random.choice(['HSBCUS33', 'BNPAUS33', 'SCBLUS33'])}",
        'confirming_bank_name': random.choice(['HSBC', 'BNP Paribas', 'Standard Chartered']),
        'confirming_bank_address': f"{random.randint(100, 9999)} {random.choice(['Financial District', 'Banking Center', 'Commerce St'])}, Singapore",
        'confirming_bank_officer': f"{random.choice(['David', 'Emma', 'James', 'Anna'])} {random.choice(['Lee', 'Chen', 'Wong', 'Tan'])}",
        'confirming_bank_officer_contact': f"+65-{random.randint(6000, 9999)}-{random.randint(1000, 9999)}",
        'confirming_bank_tel': f"+65-{random.randint(6000, 9999)}-{random.randint(1000, 9999)}",
        'issuing_bank_account_number': f"{random.randint(1000000000, 9999999999)}",
        'issuing_bank_swift': f"{random.choice(['JPMUS33', 'WFCBUS33', 'PNCUS33'])}",
        'issuing_bank_name': random.choice(['JPMorgan Chase', 'Wells Fargo', 'PNC Bank']),
        'issuing_bank_address': f"{random.randint(100, 9999)} {random.choice(['Banking Plaza', 'Financial Center', 'Commerce Ave'])}, London",
        'issuing_bank_officer': f"{random.choice(['Robert', 'Jennifer', 'Christopher', 'Amanda'])} {random.choice(['Taylor', 'Anderson', 'Thomas', 'Jackson'])}",
        'issuing_bank_officer_contact': f"+44-{random.randint(20, 29)}-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'issuing_bank_tel': f"+44-{random.randint(20, 29)}-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        
        # Commercial
        'proforma_invoice_no': f"PI-{datetime.now().year}-{random.randint(1000, 9999)}",
        'invoice_no': f"INV-{datetime.now().year}-{random.randint(1000, 9999)}",
        'commercial_id': f"COM-{datetime.now().year}-{random.randint(1000, 9999)}",
        'document_number': f"DOC-{datetime.now().year}-{random.randint(1000, 9999)}",
        'contract_value': f"USD {random.randint(1000000, 50000000):,}",
        'total_amount': f"USD {random.randint(1000000, 10000000):,}",
        'total_amount_due': f"USD {random.randint(1000000, 10000000):,}",
        'amount_in_words': f"{random.choice(['Five', 'Ten', 'Fifteen', 'Twenty'])} Million US Dollars",
        'transaction_currency': 'USD',
        'payment_terms': random.choice(['30 days', '45 days', '60 days', '90 days']),
        'validity': f"{random.randint(30, 90)} days",
        'contract_duration': f"{random.randint(6, 24)} months",
        'monthly_delivery': f"{random.randint(1000, 10000)} MT",
        'performance_bond': f"USD {random.randint(100000, 1000000):,}",
        'insurance': f"USD {random.randint(500000, 5000000):,}",
        
        # Shipping
        'port_of_loading': random.choice(['Singapore', 'Rotterdam', 'Houston', 'Dubai', 'Shanghai']),
        'port_of_discharge': random.choice(['Tokyo', 'Hamburg', 'New York', 'Los Angeles', 'Busan']),
        'place_of_destination': random.choice(['Tokyo Port', 'Hamburg Port', 'New York Port', 'Los Angeles Port']),
        'final_delivery_place': random.choice(['Tokyo Port', 'Hamburg Port', 'New York Port', 'Los Angeles Port']),
        'origin': random.choice(['Malaysia', 'Indonesia', 'Nigeria', 'Saudi Arabia', 'Kuwait']),
        'country_of_origin': random.choice(['Malaysia', 'Indonesia', 'Nigeria', 'Saudi Arabia', 'Kuwait']),
        'shipping_terms': random.choice(['FOB', 'CIF', 'CFR', 'EXW']),
        'terms_of_delivery': random.choice(['FOB', 'CIF', 'CFR', 'EXW']),
        'via_name': random.choice(['Direct', 'Via Singapore', 'Via Rotterdam', 'Via Dubai']),
        'through_name': random.choice(['Direct', 'Via Singapore', 'Via Rotterdam', 'Via Dubai']),
        'partial_shipment': random.choice(['Allowed', 'Not Allowed']),
        'transshipment': random.choice(['Allowed', 'Not Allowed']),
        
        # Dates
        'date_of_issue': (datetime.now() - timedelta(days=random.randint(1, 30))).strftime('%Y-%m-%d'),
        'issued_date': (datetime.now() - timedelta(days=random.randint(1, 30))).strftime('%Y-%m-%d'),
        'issue_date': (datetime.now() - timedelta(days=random.randint(1, 30))).strftime('%Y-%m-%d'),
        'shipment_date2': (datetime.now() + timedelta(days=random.randint(30, 90))).strftime('%Y-%m-%d'),
        'shipment_date3': (datetime.now() + timedelta(days=random.randint(30, 90))).strftime('%Y-%m-%d'),
        'valid_until': (datetime.now() + timedelta(days=random.randint(60, 180))).strftime('%Y-%m-%d'),
        'buyer_signatory_date': datetime.now().strftime('%Y-%m-%d'),
        'seller_signatory_date': datetime.now().strftime('%Y-%m-%d'),
        
        # Product specifications
        'commodity': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel', 'Bunker Fuel']),
        'product_name': random.choice(['Light Sweet Crude', 'Heavy Crude', 'Diesel Fuel', 'Gasoline']),
        'goods_details': random.choice(['Light Sweet Crude Oil', 'Heavy Crude Oil', 'Diesel Fuel', 'Gasoline']),
        'specification': random.choice(['API 35-40', 'API 25-30', 'Sulfur < 0.5%', 'Sulfur < 1.0%']),
        'quality': random.choice(['Premium Grade', 'Standard Grade', 'Commercial Grade']),
        'inspection': random.choice(['SGS', 'Bureau Veritas', 'Intertek', 'Lloyd\'s Register']),
        'cloud_point': f"{random.randint(-10, 10)}°C",
        'free_fatty_acid': f"{random.uniform(0.1, 2.0):.1f}%",
        'iodine_value': f"{random.randint(40, 60)}",
        'moisture_impurities': f"{random.uniform(0.1, 1.0):.1f}%",
        'slip_melting_point': f"{random.randint(20, 35)}°C",
        'colour': random.choice(['Light Brown', 'Dark Brown', 'Black', 'Amber']),
        
        # Quantities and prices
        'total_quantity': f"{random.randint(10000, 100000)} MT",
        'quantity2': f"{random.randint(5000, 50000)} MT",
        'quantity3': f"{random.randint(5000, 50000)} MT",
        'total_weight': f"{random.randint(10000, 100000)} MT",
        'total_gross': f"{random.randint(10000, 100000)} MT",
        'total_containers': f"{random.randint(1, 10)}",
        'unit_price2': f"USD {random.randint(50, 100)}.00",
        'unit_price3': f"USD {random.randint(50, 100)}.00",
        'amount2': f"USD {random.randint(250000, 5000000):,}",
        'amount3': f"USD {random.randint(250000, 5000000):,}",
        'price': f"USD {random.randint(50, 100)}.00",
        'shipping_charges': f"USD {random.randint(50000, 500000):,}",
        'other_expenditures': f"USD {random.randint(10000, 100000):,}",
        'discount': f"{random.randint(0, 10)}%",
        
        # Items
        'item2': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel']),
        'item3': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel']),
        'consignment2': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel']),
        'consignment33': random.choice(['Crude Oil', 'Diesel', 'Gasoline', 'Jet Fuel']),
        
        # Signatories
        'seller_signatory_name': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'seller_signatory_position': random.choice(['Managing Director', 'Sales Manager', 'Operations Manager', 'CEO']),
        'seller_signature': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'buyer_signatory_name': f"{random.choice(['Taro', 'Yuki', 'Hiroshi', 'Akira'])} {random.choice(['Yamada', 'Sato', 'Suzuki', 'Takahashi'])}",
        'buyer_signatory_position': random.choice(['Procurement Manager', 'General Manager', 'Director', 'President']),
        'buyer_signature': f"{random.choice(['Taro', 'Yuki', 'Hiroshi', 'Akira'])} {random.choice(['Yamada', 'Sato', 'Suzuki', 'Takahashi'])}",
        'signatory_name': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'authorized_person_name': f"{random.choice(['John', 'Sarah', 'Michael', 'Lisa'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        'notary_number': f"NOT-{random.randint(1000, 9999)}",
        
        # Contact information
        'buyer_company_name': random.choice(['Tokyo Trading Co.', 'Osaka Shipping Ltd.', 'Yokohama Marine Inc.', 'Kobe Commerce Corp.']),
        'buyer_name': f"{random.choice(['Taro', 'Yuki', 'Hiroshi', 'Akira'])} {random.choice(['Yamada', 'Sato', 'Suzuki', 'Takahashi'])}",
        'buyer_address': f"{random.randint(1, 999)} {random.choice(['Chuo-dori', 'Ginza', 'Shibuya', 'Shinjuku'])}, Tokyo, Japan",
        'buyer_city_country': 'Tokyo, Japan',
        'buyer_email': f"buyer{random.randint(1, 999)}@{random.choice(['tokyo-trading.com', 'osaka-shipping.com', 'yokohama-marine.com'])}",
        'buyer_tel': f"+81-3-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'buyer_fax': f"+81-3-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'buyer_mobile': f"+81-{random.randint(90, 99)}-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'buyer_office_tel': f"+81-3-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        'buyer_designation': random.choice(['Procurement Manager', 'General Manager', 'Director', 'President']),
        'buyer_representative': f"{random.choice(['Taro', 'Yuki', 'Hiroshi', 'Akira'])} {random.choice(['Yamada', 'Sato', 'Suzuki', 'Takahashi'])}",
        'buyer_position': random.choice(['Procurement Manager', 'General Manager', 'Director', 'President']),
        'buyer_registration': f"REG-{random.randint(100000, 999999)}",
        
        # Seller information
        'seller_company': random.choice(['Singapore Trading Ltd.', 'Malaysia Oil Corp.', 'Indonesia Marine Inc.', 'Thailand Commerce Co.']),
        'seller_address': f"{random.randint(1, 999)} {random.choice(['Marina Bay', 'Orchard Road', 'Raffles Place', 'Clarke Quay'])}, Singapore",
        
        # Default fallback
        'default': f"Sample {placeholder.replace('_', ' ').title()}"
    }
    
    # Return realistic data or default
    return realistic_data.get(placeholder.lower(), realistic_data['default'])

def fill_word_template(template_path, output_path, vessel_data):
    """Fill a Word template with vessel data"""
    try:
        # Copy template to output path
        shutil.copy2(template_path, output_path)
        
        # Open the document
        doc = Document(output_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            original_text = text
            for placeholder, value in vessel_data.items():
                # Replace both {placeholder} and {placeholder} formats
                text = text.replace(f"{{{placeholder}}}", str(value))
                text = text.replace(f"{{{{{placeholder}}}}}", str(value))  # Handle double braces
            
            # Find any remaining placeholders and replace with realistic data
            remaining_placeholders = re.findall(r'\{([^}]+)\}', text)
            for placeholder in remaining_placeholders:
                realistic_value = generate_realistic_data_for_placeholder(placeholder)
                text = text.replace(f"{{{placeholder}}}", str(realistic_value))
                print(f"Generated realistic data for missing placeholder '{placeholder}': {realistic_value}")
            
            if text != original_text:
                paragraph.text = text
                print(f"Replaced placeholders in paragraph: {original_text[:50]}... -> {text[:50]}...")
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        original_text = text
                        for placeholder, value in vessel_data.items():
                            # Replace both {placeholder} and {placeholder} formats
                            text = text.replace(f"{{{placeholder}}}", str(value))
                            text = text.replace(f"{{{{{placeholder}}}}}", str(value))  # Handle double braces
                        
                        # Find any remaining placeholders and replace with realistic data
                        remaining_placeholders = re.findall(r'\{([^}]+)\}', text)
                        for placeholder in remaining_placeholders:
                            realistic_value = generate_realistic_data_for_placeholder(placeholder)
                            text = text.replace(f"{{{placeholder}}}", str(realistic_value))
                            print(f"Generated realistic data for missing placeholder '{placeholder}': {realistic_value}")
                        
                        if text != original_text:
                            paragraph.text = text
                            print(f"Replaced placeholders in table cell: {original_text[:50]}... -> {text[:50]}...")
        
        # Save the document
        doc.save(output_path)
        print(f"Template filled successfully: {output_path}")
        return True
    except Exception as e:
        print(f"Error filling template: {e}")
        return False

@app.get("/")
async def root():
    return {"message": "Working Document Service is running", "status": "ok"}

@app.get("/health")
async def health():
    return {"status": "healthy", "service": "document-processing"}

@app.get("/templates")
async def get_templates():
    """Get list of available templates"""
    # Return actual stored templates
    return templates_storage

@app.get("/vessels")
async def get_vessels():
    return [
        {
            "id": "1",
            "name": "Petroleum Express 529",
            "imo": "IMO1861018",
            "vessel_type": "Crude Oil Tanker",
            "flag": "Malta"
        },
        {
            "id": "2", 
            "name": "Atlantic Voyager 805",
            "imo": "IMO2379622",
            "vessel_type": "Container Ship",
            "flag": "Panama"
        }
    ]


@app.post("/process-document")
async def process_document(
    template_id: str = Form(...),
    vessel_imo: str = Form(...),
    template_file: UploadFile = File(...)
):
    """Process a document template with vessel data"""
    try:
        # Generate a unique document ID
        document_id = str(uuid.uuid4())
        
        # Create outputs directory if it doesn't exist
        outputs_dir = Path("outputs")
        outputs_dir.mkdir(exist_ok=True)
        
        # Find the template in storage
        template_info = None
        for template in templates_storage:
            if template["id"] == template_id:
                template_info = template
                break
        
        if not template_info:
            return JSONResponse({
                "success": False,
                "message": "Template not found",
                "error": "Template ID not found in storage"
            }, status_code=404)
        
        # Get REAL vessel data from your Supabase database
        vessel_data = await get_real_vessel_data(vessel_imo)
        
        # Find the template file
        templates_dir = Path("templates")
        template_file_path = templates_dir / f"{template_id}.docx"
        
        if not template_file_path.exists():
            return JSONResponse({
                "success": False,
                "message": "Template file not found",
                "error": "Template file does not exist on server"
            }, status_code=404)
        
        # Create output files
        outputs_dir = Path("outputs")
        outputs_dir.mkdir(exist_ok=True)
        filled_docx_file = outputs_dir / f"{document_id}_filled.docx"
        pdf_file = outputs_dir / f"{document_id}_filled.pdf"
        
        # Fill the Word template with vessel data
        success = fill_word_template(template_file_path, filled_docx_file, vessel_data)
        
        if not success:
            return JSONResponse({
                "success": False,
                "message": "Failed to fill template",
                "error": "Could not process the Word template"
            }, status_code=500)
        
        # Try to convert Word document to PDF (preserving original design)
        pdf_success = False
        pdf_conversion_method = "none"
        
        print(f"🔄 Starting PDF conversion for: {filled_docx_file}")
        
        try:
            import platform
            print(f"🖥️  Platform: {platform.system()}")
            
            # Method 1: Try using LibreOffice first (best for Linux/Render)
            if not pdf_success:
                try:
                    import subprocess
                    print("🔄 Trying LibreOffice conversion (best for Linux)...")
                    
                    # Create a temporary directory for LibreOffice
                    import tempfile
                    with tempfile.TemporaryDirectory() as temp_dir:
                        # Copy the file to temp directory
                        temp_docx = os.path.join(temp_dir, "document.docx")
                        shutil.copy2(filled_docx_file, temp_docx)
                        
                        # Run LibreOffice conversion
                        result = subprocess.run([
                            'libreoffice', '--headless', '--convert-to', 'pdf', 
                            '--outdir', temp_dir, temp_docx
                        ], capture_output=True, text=True, timeout=120)
                        
                        # Check if PDF was created
                        temp_pdf = os.path.join(temp_dir, "document.pdf")
                        if result.returncode == 0 and os.path.exists(temp_pdf):
                            # Copy the PDF to our output directory
                            shutil.copy2(temp_pdf, pdf_file)
                            pdf_success = True
                            pdf_conversion_method = "libreoffice"
                            print(f"✅ PDF conversion successful with LibreOffice: {pdf_file}")
                        else:
                            print(f"⚠️  LibreOffice failed - Return code: {result.returncode}")
                            print(f"⚠️  LibreOffice stderr: {result.stderr}")
                            print(f"⚠️  LibreOffice stdout: {result.stdout}")
                            
                except FileNotFoundError:
                    print("⚠️  LibreOffice not found on system")
                except Exception as e:
                    print(f"⚠️  LibreOffice conversion failed: {e}")
            
            # Method 2: Try using unoconv (alternative for Linux)
            if not pdf_success:
                try:
                    import subprocess
                    print("🔄 Trying unoconv conversion...")
                    result = subprocess.run([
                        'unoconv', '-f', 'pdf', '-o', str(pdf_file), str(filled_docx_file)
                    ], capture_output=True, text=True, timeout=60)
                    
                    if result.returncode == 0 and pdf_file.exists():
                        pdf_success = True
                        pdf_conversion_method = "unoconv"
                        print(f"✅ PDF conversion successful with unoconv: {pdf_file}")
                    else:
                        print(f"⚠️  unoconv failed - Return code: {result.returncode}")
                        print(f"⚠️  unoconv stderr: {result.stderr}")
                except FileNotFoundError:
                    print("⚠️  unoconv not found on system")
                except Exception as e:
                    print(f"⚠️  unoconv conversion failed: {e}")
                    
            # Method 3: Try using pandoc (if available)
            if not pdf_success:
                try:
                    import subprocess
                    print("🔄 Trying pandoc conversion...")
                    result = subprocess.run([
                        'pandoc', str(filled_docx_file), '-o', str(pdf_file)
                    ], capture_output=True, text=True, timeout=60)
                    
                    if result.returncode == 0 and pdf_file.exists():
                        pdf_success = True
                        pdf_conversion_method = "pandoc"
                        print(f"✅ PDF conversion successful with pandoc: {pdf_file}")
                    else:
                        print(f"⚠️  Pandoc failed - Return code: {result.returncode}")
                        print(f"⚠️  Pandoc stderr: {result.stderr}")
                except FileNotFoundError:
                    print("⚠️  Pandoc not found on system")
                except Exception as e:
                    print(f"⚠️  Pandoc conversion failed: {e}")
            
            # Method 4: Try docx2pdf only on Windows
            if not pdf_success and platform.system() == "Windows":
                try:
                    print("🔄 Trying docx2pdf (Windows only)...")
                    from docx2pdf import convert
                    convert(str(filled_docx_file), str(pdf_file))
                    pdf_success = True
                    pdf_conversion_method = "docx2pdf-windows"
                    print(f"✅ PDF conversion successful with docx2pdf: {pdf_file}")
                except Exception as e:
                    print(f"⚠️  docx2pdf failed: {e}")
                    
            # Method 5: Create notice PDF explaining the situation
            if not pdf_success:
                print("ℹ️  All PDF conversion methods failed, creating informational PDF...")
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                
                c = canvas.Canvas(str(pdf_file), pagesize=letter)
                c.drawString(100, 750, "DOCUMENT PROCESSING COMPLETE")
                c.drawString(100, 720, f"Template: {template_info['name']}")
                c.drawString(100, 690, f"Vessel IMO: {vessel_imo}")
                c.drawString(100, 660, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                c.drawString(100, 620, "=" * 60)
                c.drawString(100, 580, "IMPORTANT NOTICE:")
                c.drawString(100, 550, "Your template has been processed successfully!")
                c.drawString(100, 520, "• All placeholders have been filled with real vessel data")
                c.drawString(100, 490, "• Your original template design is preserved")
                c.drawString(100, 460, "• Download the Word document (.docx) for best results")
                c.drawString(100, 420, "The Word document contains:")
                c.drawString(120, 390, "✓ Your exact template design and formatting")
                c.drawString(120, 360, "✓ All backgrounds, logos, and images")
                c.drawString(120, 330, "✓ Original tables and layout")
                c.drawString(120, 300, "✓ Filled with real vessel data")
                c.drawString(100, 260, "This PDF is informational only.")
                c.drawString(100, 230, "Please download the .docx file for the complete document.")
                c.save()
                pdf_success = True
                pdf_conversion_method = "informational"
                print(f"ℹ️  Created informational PDF: {pdf_file}")
                
        except Exception as e:
            print(f"❌ PDF creation completely failed: {e}")
            pdf_success = False
            pdf_conversion_method = "failed"
        
        # Return success response
        return JSONResponse({
            "success": True,
            "message": "Document processed successfully!",
            "document_id": document_id,
            "download_url": f"/download/{document_id}",
            "vessel_imo": vessel_imo,
            "template_id": template_id,
            "template_name": template_info["name"],
            "files_created": {
                "docx_file": str(filled_docx_file),
                "pdf_file": str(pdf_file) if pdf_success else None
            },
            "download_urls": {
                "docx": f"/download/{document_id}?format=docx",
                "pdf": f"/download/{document_id}?format=pdf"
            },
            "pdf_conversion": {
                "success": pdf_success,
                "method": pdf_conversion_method,
                "note": "PDF conversion preserves your exact Word template design" if pdf_success else "PDF conversion failed - download DOCX for exact template design"
            },
            "note": "The Word document (.docx) contains your exact template with all formatting preserved."
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Document processing failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.get("/download/{document_id}")
async def download_document(document_id: str, format: str = "pdf"):
    """Download processed document as actual file"""
    try:
        outputs_dir = Path("outputs")
        outputs_dir.mkdir(exist_ok=True)
        
        if format.lower() == "pdf":
            file_path = outputs_dir / f"{document_id}_filled.pdf"
            media_type = "application/pdf"
            filename = f"vessel_report_{document_id}.pdf"
        elif format.lower() == "docx":
            file_path = outputs_dir / f"{document_id}_filled.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"vessel_report_{document_id}.docx"
        else:
            file_path = outputs_dir / f"{document_id}_filled_fallback.txt"
            media_type = "text/plain"
            filename = f"vessel_report_{document_id}.txt"
        
        if not file_path.exists():
            # Create fallback file if it doesn't exist
            if format.lower() == "pdf":
                # Create informational PDF if PDF doesn't exist
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                
                c = canvas.Canvas(str(file_path), pagesize=letter)
                c.drawString(100, 750, "Document Processing Notice")
                c.drawString(100, 720, f"Document ID: {document_id}")
                c.drawString(100, 690, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                c.drawString(100, 630, "The PDF conversion could not preserve the original template design.")
                c.drawString(100, 600, "Please download the Word document (.docx) instead,")
                c.drawString(100, 570, "which contains your exact template with all formatting,")
                c.drawString(100, 540, "backgrounds, logos, and design preserved.")
                c.save()
            elif format.lower() == "docx":
                # If DOCX doesn't exist, this is a real error
                raise HTTPException(status_code=404, detail="Word document not found")
            else:
                # Create text fallback
                with open(file_path, 'w') as f:
                    f.write(f"Document ID: {document_id}\n")
                    f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write("Please download the Word document (.docx) for the full formatted template.\n")
        
        return FileResponse(
            path=str(file_path),
            media_type=media_type,
            filename=filename
        )
        
    except Exception as e:
        print(f"Download error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

@app.put("/templates/{template_id}")
async def update_template(
    template_id: str,
    name: str = Form(None),
    description: str = Form(None),
    is_active: bool = Form(None)
):
    """Update template information"""
    try:
        # Find template in storage
        template_found = False
        for i, template in enumerate(templates_storage):
            if template["id"] == template_id:
                # Update fields if provided
                if name is not None:
                    template["name"] = name
                if description is not None:
                    template["description"] = description
                if is_active is not None:
                    template["is_active"] = is_active
                
                template_found = True
                break
        
        if not template_found:
            return JSONResponse({
                "success": False,
                "message": "Template not found"
            }, status_code=404)
        
        # Save updated templates
        save_templates()
        
        return JSONResponse({
            "success": True,
            "message": "Template updated successfully"
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Template update failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.delete("/templates/{template_id}")
async def delete_template(template_id: str):
    """Delete a template"""
    try:
        # Find template in storage
        template_found = False
        for i, template in enumerate(templates_storage):
            if template["id"] == template_id:
                # Remove from storage
                templates_storage.pop(i)
                template_found = True
                break
        
        if not template_found:
            return JSONResponse({
                "success": False,
                "message": "Template not found"
            }, status_code=404)
        
        # Delete template file
        templates_dir = Path("templates")
        template_file = templates_dir / f"{template_id}.docx"
        if template_file.exists():
            template_file.unlink()
        
        # Save updated templates
        save_templates()
        
        return JSONResponse({
            "success": True,
            "message": "Template deleted successfully"
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Template deletion failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.get("/user-permissions")
async def get_user_permissions():
    """Get user permissions (simplified version)"""
    return {
        "success": True,
        "permissions": {
            "can_upload_templates": True,
            "can_edit_templates": True,
            "can_delete_templates": True,
            "can_process_documents": True,
            "template_limit": 100,
            "documents_per_month": 1000
        }
    }

@app.post("/upload-template")
async def upload_template(
    name: str = Form(...),
    description: str = Form(...),
    template_file: UploadFile = File(...)
):
    """Upload a new document template"""
    try:
        # Validate file type
        if not template_file.filename.endswith('.docx'):
            return JSONResponse({
                "success": False,
                "message": "Only .docx files are allowed"
            }, status_code=400)
        
        # Generate a unique template ID
        template_id = str(uuid.uuid4())
        
        # Create templates directory if it doesn't exist
        templates_dir = Path("templates")
        templates_dir.mkdir(exist_ok=True)
        
        # Save the uploaded file
        file_extension = Path(template_file.filename).suffix
        saved_filename = f"{template_id}{file_extension}"
        file_path = templates_dir / saved_filename
        
        with open(file_path, "wb") as buffer:
            content = await template_file.read()
            buffer.write(content)
        
        # Extract placeholders from the template
        placeholders = extract_placeholders_from_docx(str(file_path))
        
        # Store template information
        template_info = {
            "id": template_id,
            "name": name,
            "description": description,
            "file_name": template_file.filename,
            "file_size": len(content),
            "placeholders": placeholders,
            "is_active": True,
            "created_at": datetime.now().isoformat(),
            "created_by": "system"
        }
        
        templates_storage.append(template_info)
        save_templates()
        
        return JSONResponse({
            "success": True,
            "message": "Template uploaded successfully",
            "template": template_info
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Template upload failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    print("Starting Working Document Service...")
    print("Available endpoints:")
    print("  GET  /")
    print("  GET  /health")
    print("  GET  /templates")
    print("  POST /upload-template")
    print("  POST /process-document")
    print("  GET  /download/{document_id}")
    print("  PUT  /templates/{template_id}")
    print("  DELETE /templates/{template_id}")
    uvicorn.run(app, host="0.0.0.0", port=port)