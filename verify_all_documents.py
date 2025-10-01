#!/usr/bin/env python3
"""
Verification Script for All Processed Documents
Checks the quality of filled documents and provides detailed analysis
"""

import os
import glob
from docx import Document
import re

def analyze_document(file_path):
    """Analyze a single document for placeholder filling quality"""
    try:
        doc = Document(file_path)
        
        # Extract all text from the document
        total_text = ""
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            total_text += paragraph.text + "\n"
        
        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_text += paragraph.text + "\n"
        
        # Count different types of placeholders and issues
        analysis = {
            'total_chars': len(total_text),
            'unfilled_curly_single': len(re.findall(r'\{[^}]*\}', total_text)),
            'unfilled_curly_double': len(re.findall(r'\{\{[^}]*\}\}', total_text)),
            'unfilled_square_single': len(re.findall(r'\[[^\]]*\]', total_text)) - total_text.count('[Data Not Available]'),
            'unfilled_square_double': len(re.findall(r'\[\[[^\]]*\]\]', total_text)),
            'data_not_available': total_text.count('[Data Not Available]'),
            'empty_curly': total_text.count('{}') + total_text.count('{{}}'),
            'empty_square': total_text.count('[]') + total_text.count('[[]]'),
            'malformed_curly': len(re.findall(r'\{[^}]*$|^[^{]*\}', total_text)),
            'malformed_square': len(re.findall(r'\[[^\]]*$|^[^\[]*\]', total_text)),
        }
        
        # Calculate quality score
        total_issues = (analysis['unfilled_curly_single'] + 
                       analysis['unfilled_curly_double'] + 
                       analysis['unfilled_square_single'] + 
                       analysis['unfilled_square_double'] + 
                       analysis['data_not_available'] + 
                       analysis['empty_curly'] + 
                       analysis['empty_square'] + 
                       analysis['malformed_curly'] + 
                       analysis['malformed_square'])
        
        analysis['total_issues'] = total_issues
        analysis['quality_score'] = max(0, 100 - (total_issues * 2))  # Deduct 2 points per issue
        
        # Sample some filled content to verify realistic data
        sample_lines = [line.strip() for line in total_text.split('\n') if line.strip()][:10]
        analysis['sample_content'] = sample_lines
        
        return analysis
        
    except Exception as e:
        return {'error': str(e)}

def get_original_name(filename):
    """Extract original document name from processed filename"""
    # Remove the UUID prefix and _filled.docx suffix
    name = filename.replace('_filled.docx', '')
    if len(name) > 36 and name[8] == '-':
        return name[37:]  # Remove UUID prefix
    return name

def main():
    """Main verification function"""
    print("🔍 DOCUMENT VERIFICATION REPORT")
    print("=" * 60)
    
    # Find all processed documents
    outputs_dir = "outputs"
    if not os.path.exists(outputs_dir):
        print("❌ Outputs directory not found!")
        return
    
    filled_docs = glob.glob(os.path.join(outputs_dir, "*_filled.docx"))
    
    if not filled_docs:
        print("❌ No filled documents found!")
        return
    
    print(f"📋 Found {len(filled_docs)} processed documents to verify")
    print()
    
    # Analyze each document
    results = []
    perfect_docs = 0
    good_docs = 0
    needs_attention = 0
    
    for file_path in filled_docs:
        filename = os.path.basename(file_path)
        original_name = get_original_name(filename)
        
        print(f"🔍 Analyzing: {original_name}")
        
        analysis = analyze_document(file_path)
        
        if 'error' in analysis:
            print(f"    ❌ Error: {analysis['error']}")
            continue
        
        results.append({
            'filename': filename,
            'original_name': original_name,
            'analysis': analysis
        })
        
        # Categorize document quality
        if analysis['total_issues'] == 0:
            perfect_docs += 1
            print(f"    ✅ Perfect! No issues found")
        elif analysis['quality_score'] >= 80:
            good_docs += 1
            print(f"    ✅ Good quality (Score: {analysis['quality_score']}/100)")
        else:
            needs_attention += 1
            print(f"    ⚠️  Needs attention (Score: {analysis['quality_score']}/100)")
        
        # Show detailed issues if any
        if analysis['total_issues'] > 0:
            issues = []
            if analysis['unfilled_curly_single'] > 0:
                issues.append(f"Single curly braces: {analysis['unfilled_curly_single']}")
            if analysis['unfilled_curly_double'] > 0:
                issues.append(f"Double curly braces: {analysis['unfilled_curly_double']}")
            if analysis['unfilled_square_single'] > 0:
                issues.append(f"Single square brackets: {analysis['unfilled_square_single']}")
            if analysis['unfilled_square_double'] > 0:
                issues.append(f"Double square brackets: {analysis['unfilled_square_double']}")
            if analysis['data_not_available'] > 0:
                issues.append(f"Data not available: {analysis['data_not_available']}")
            if analysis['empty_curly'] > 0:
                issues.append(f"Empty curly braces: {analysis['empty_curly']}")
            if analysis['empty_square'] > 0:
                issues.append(f"Empty square brackets: {analysis['empty_square']}")
            
            print(f"    📊 Issues: {', '.join(issues)}")
        
        print(f"    📄 Document size: {analysis['total_chars']:,} characters")
        print()
    
    # Generate summary
    print("📊 VERIFICATION SUMMARY")
    print("=" * 60)
    print(f"📄 Total documents verified: {len(results)}")
    print(f"✅ Perfect documents: {perfect_docs}")
    print(f"👍 Good quality documents: {good_docs}")
    print(f"⚠️  Documents needing attention: {needs_attention}")
    print()
    
    if perfect_docs > 0:
        print("🎉 PERFECT DOCUMENTS (No issues):")
        for result in results:
            if result['analysis']['total_issues'] == 0:
                print(f"  ✅ {result['original_name']}")
        print()
    
    if good_docs > 0:
        print("👍 GOOD QUALITY DOCUMENTS (Minor issues):")
        for result in results:
            analysis = result['analysis']
            if 0 < analysis['total_issues'] <= 10 and analysis['quality_score'] >= 80:
                print(f"  ✅ {result['original_name']} (Score: {analysis['quality_score']}/100)")
        print()
    
    if needs_attention > 0:
        print("⚠️  DOCUMENTS NEEDING ATTENTION:")
        for result in results:
            analysis = result['analysis']
            if analysis['quality_score'] < 80:
                print(f"  ⚠️  {result['original_name']} (Score: {analysis['quality_score']}/100)")
                print(f"      Issues: {analysis['total_issues']} total")
                if analysis['unfilled_curly_single'] > 0:
                    print(f"      - Unfilled single curly braces: {analysis['unfilled_curly_single']}")
                if analysis['unfilled_square_single'] > 0:
                    print(f"      - Unfilled single square brackets: {analysis['unfilled_square_single']}")
                if analysis['data_not_available'] > 0:
                    print(f"      - Data not available markers: {analysis['data_not_available']}")
        print()
    
    # Overall quality assessment
    total_docs = len(results)
    if total_docs > 0:
        success_rate = ((perfect_docs + good_docs) / total_docs) * 100
        print(f"📈 Overall Success Rate: {success_rate:.1f}%")
        
        if success_rate >= 90:
            print("🎉 Excellent! Most documents are perfectly filled!")
        elif success_rate >= 70:
            print("👍 Good! Most documents are well filled with minor issues.")
        else:
            print("⚠️  Some documents need attention for better placeholder filling.")
    
    print()
    print("✅ Verification completed!")

if __name__ == "__main__":
    main()