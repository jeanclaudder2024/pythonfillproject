#!/usr/bin/env python3
"""
Analyze Unfilled Placeholders Script
Identifies specific patterns of unfilled placeholders in processed documents
"""

import os
import glob
from docx import Document
import re
from collections import Counter

def extract_unfilled_placeholders(file_path):
    """Extract all unfilled placeholders from a document"""
    try:
        doc = Document(file_path)
        
        # Extract all text from the document
        total_text = ""
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            total_text += paragraph.text + "\n"
        
        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_text += paragraph.text + "\n"
        
        # Find all unfilled placeholders
        unfilled = {
            'curly_single': re.findall(r'\{[^}]*\}', total_text),
            'curly_double': re.findall(r'\{\{[^}]*\}\}', total_text),
            'square_single': [match for match in re.findall(r'\[[^\]]*\]', total_text) if match != '[Data Not Available]'],
            'square_double': re.findall(r'\[\[[^\]]*\]\]', total_text),
        }
        
        return unfilled
        
    except Exception as e:
        return {'error': str(e)}

def main():
    """Main analysis function"""
    print("🔍 ANALYZING UNFILLED PLACEHOLDERS")
    print("=" * 60)
    
    # Find all processed documents
    outputs_dir = "outputs"
    filled_docs = glob.glob(os.path.join(outputs_dir, "*_filled.docx"))
    
    if not filled_docs:
        print("❌ No filled documents found!")
        return
    
    print(f"📋 Analyzing {len(filled_docs)} processed documents")
    print()
    
    # Collect all unfilled placeholders
    all_unfilled = {
        'curly_single': [],
        'curly_double': [],
        'square_single': [],
        'square_double': []
    }
    
    problem_docs = []
    
    for file_path in filled_docs:
        filename = os.path.basename(file_path)
        
        unfilled = extract_unfilled_placeholders(file_path)
        
        if 'error' in unfilled:
            print(f"❌ Error analyzing {filename}: {unfilled['error']}")
            continue
        
        # Count total unfilled
        total_unfilled = (len(unfilled['curly_single']) + 
                         len(unfilled['curly_double']) + 
                         len(unfilled['square_single']) + 
                         len(unfilled['square_double']))
        
        if total_unfilled > 0:
            problem_docs.append({
                'filename': filename,
                'unfilled': unfilled,
                'total': total_unfilled
            })
            
            # Add to global collection
            for key in all_unfilled:
                all_unfilled[key].extend(unfilled[key])
    
    # Analyze patterns
    print("📊 UNFILLED PLACEHOLDER PATTERNS")
    print("=" * 60)
    
    if all_unfilled['curly_single']:
        print("🔸 UNFILLED SINGLE CURLY BRACES:")
        counter = Counter(all_unfilled['curly_single'])
        for placeholder, count in counter.most_common(10):
            print(f"  {placeholder} (appears {count} times)")
        print()
    
    if all_unfilled['square_single']:
        print("🔸 UNFILLED SINGLE SQUARE BRACKETS:")
        counter = Counter(all_unfilled['square_single'])
        for placeholder, count in counter.most_common(20):
            print(f"  {placeholder} (appears {count} times)")
        print()
    
    if all_unfilled['curly_double']:
        print("🔸 UNFILLED DOUBLE CURLY BRACES:")
        counter = Counter(all_unfilled['curly_double'])
        for placeholder, count in counter.most_common(10):
            print(f"  {placeholder} (appears {count} times)")
        print()
    
    if all_unfilled['square_double']:
        print("🔸 UNFILLED DOUBLE SQUARE BRACKETS:")
        counter = Counter(all_unfilled['square_double'])
        for placeholder, count in counter.most_common(10):
            print(f"  {placeholder} (appears {count} times)")
        print()
    
    # Show problem documents
    if problem_docs:
        print("📄 DOCUMENTS WITH UNFILLED PLACEHOLDERS:")
        print("=" * 60)
        
        for doc in sorted(problem_docs, key=lambda x: x['total'], reverse=True):
            print(f"📄 {doc['filename']} ({doc['total']} unfilled)")
            
            if doc['unfilled']['curly_single']:
                print(f"  🔸 Single curly: {len(doc['unfilled']['curly_single'])} - {doc['unfilled']['curly_single'][:3]}...")
            if doc['unfilled']['square_single']:
                print(f"  🔸 Single square: {len(doc['unfilled']['square_single'])} - {doc['unfilled']['square_single'][:3]}...")
            if doc['unfilled']['curly_double']:
                print(f"  🔸 Double curly: {len(doc['unfilled']['curly_double'])} - {doc['unfilled']['curly_double'][:3]}...")
            if doc['unfilled']['square_double']:
                print(f"  🔸 Double square: {len(doc['unfilled']['square_double'])} - {doc['unfilled']['square_double'][:3]}...")
            print()
    
    print("✅ Analysis completed!")

if __name__ == "__main__":
    main()