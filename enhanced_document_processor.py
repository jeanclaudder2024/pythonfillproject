import os
import re
import requests
from docx import Document
from bs4 import BeautifulSoup
import random
from faker import Faker
import subprocess
from datetime import datetime
from dotenv import load_dotenv
import json
from random_data_generator import RandomDataGenerator
from database_integration import SupabaseIntegration

# Load environment variables
load_dotenv()

class EnhancedDocumentProcessor:
    def __init__(self):
        self.fake = Faker()
        self.vessel_data_cache = {}
        self.random_generator = RandomDataGenerator()
        self.db_integration = SupabaseIntegration()
        print("SUCCESS: Enhanced document processor initialized with database integration")
        
    def process_document(self, input_path, vessel_imo=None, vessel_id=None, file_id=None):
        """Process the document by filling placeholders with real or random data"""
        # Load the document
        doc = Document(input_path)
        
        # Find all placeholders
        placeholders = self.find_placeholders(doc)
        print(f"Found {len(placeholders)} placeholders: {list(placeholders)}")
        
        # Get real data from database if available
        real_data = {}
        if vessel_imo:
            real_data = self.db_integration.get_vessel_data(vessel_imo)
        elif vessel_id:
            real_data = self.db_integration.get_vessel_by_id(vessel_id)
        
        # Generate replacement data (real data + fallback random data)
        replacement_data = self.generate_smart_replacement_data(placeholders, real_data, vessel_imo)
        
        # Fill placeholders
        self.fill_placeholders(doc, placeholders, replacement_data)
        
        # Save the filled Word document
        word_output = os.path.join('outputs', f"{file_id}_filled.docx")
        doc.save(word_output)
        print(f"SUCCESS: Saved filled document: {word_output}")
        
        # Convert to PDF
        pdf_output = os.path.join('outputs', f"{file_id}_filled.pdf")
        self.convert_to_pdf(word_output, pdf_output)
        
        # Create fallback text file
        fallback_output = os.path.join('outputs', f"{file_id}_filled_fallback.txt")
        self.create_fallback_text(replacement_data, fallback_output)
        
        return word_output, pdf_output
    
    def find_placeholders(self, doc):
        """Find all placeholders in the document (format: {placeholder_name} or {{placeholder_name}})"""
        placeholders = set()
        
        # Check paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Find placeholders in format {name} or {{name}}
            matches = re.findall(r'\{\{?([^}]+)\}?\}', text)
            placeholders.update(matches)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        matches = re.findall(r'\{\{?([^}]+)\}?\}', text)
                        placeholders.update(matches)
        
        # Check headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    text = paragraph.text
                    matches = re.findall(r'\{\{?([^}]+)\}?\}', text)
                    placeholders.update(matches)
            
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    text = paragraph.text
                    matches = re.findall(r'\{\{?([^}]+)\}?\}', text)
                    placeholders.update(matches)
        
        return placeholders
    
    def generate_smart_replacement_data(self, placeholders, real_data, vessel_imo=None):
        """Generate replacement data using real data when available, fallback to random data"""
        replacement_data = {}
        
        # Add current date/time data
        now = datetime.now()
        replacement_data.update({
            'date': now.strftime('%Y-%m-%d'),
            'current_date': now.strftime('%Y-%m-%d'),
            'today': now.strftime('%Y-%m-%d'),
            'time': now.strftime('%H:%M:%S'),
            'current_time': now.strftime('%H:%M:%S'),
            'year': now.strftime('%Y'),
            'month': now.strftime('%m'),
            'day': now.strftime('%d'),
        })
        
        # Use real data when available
        for placeholder in placeholders:
            if placeholder in real_data and real_data[placeholder]:
                replacement_data[placeholder] = real_data[placeholder]
                print(f"SUCCESS: Using real data for {placeholder}: {real_data[placeholder]}")
            else:
                # Generate random data as fallback
                random_value = self.random_generator.get_random_value(placeholder)
                replacement_data[placeholder] = random_value
                print(f"INFO: Using random data for {placeholder}: {random_value}")
        
        return replacement_data
    
    def fill_placeholders(self, doc, placeholders, replacement_data):
        """Fill placeholders in the document with replacement data"""
        for placeholder in placeholders:
            value = replacement_data.get(placeholder, f"[{placeholder}]")
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                self.replace_in_paragraph(paragraph, placeholder, value)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_in_paragraph(paragraph, placeholder, value)
            
            # Replace in headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        self.replace_in_paragraph(paragraph, placeholder, value)
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        self.replace_in_paragraph(paragraph, placeholder, value)
    
    def replace_in_paragraph(self, paragraph, placeholder, value):
        """Replace placeholder in a paragraph"""
        # Try both formats: {placeholder} and {{placeholder}}
        patterns = [
            f"{{{{{placeholder}}}}}",  # {{placeholder}}
            f"{{{placeholder}}}"       # {placeholder}
        ]
        
        for pattern in patterns:
            if pattern in paragraph.text:
                # Replace the text
                paragraph.text = paragraph.text.replace(pattern, str(value))
                print(f"Replaced {pattern} with {value}")
    
    def convert_to_pdf(self, word_path, pdf_path):
        """Convert Word document to PDF"""
        try:
            # Try using docx2pdf first
            from docx2pdf import convert
            convert(word_path, pdf_path)
            print(f"SUCCESS: PDF conversion successful: {pdf_path}")
        except Exception as e:
            print(f"WARNING: docx2pdf failed: {str(e)}")
            try:
                # Try using LibreOffice
                subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf', 
                    '--outdir', os.path.dirname(pdf_path), word_path
                ], check=True, capture_output=True)
                print(f"SUCCESS: PDF conversion successful with LibreOffice: {pdf_path}")
            except Exception as e2:
                print(f"ERROR: PDF conversion failed: {str(e2)}")
                # Create a text file indicating PDF conversion failed
                with open(pdf_path.replace('.pdf', '_conversion_failed.txt'), 'w') as f:
                    f.write(f"PDF conversion failed for {word_path}\n")
                    f.write(f"Error: {str(e2)}\n")
                    f.write("Please use the Word document instead.\n")
    
    def create_fallback_text(self, replacement_data, output_path):
        """Create a text file with all the replacement data"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("Document Placeholder Data\n")
                f.write("=" * 50 + "\n\n")
                
                for placeholder, value in replacement_data.items():
                    f.write(f"{placeholder}: {value}\n")
                
                f.write(f"\nGenerated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            
            print(f"SUCCESS: Created fallback text file: {output_path}")
        except Exception as e:
            print(f"ERROR: Failed to create fallback text file: {str(e)}")
    
    def get_vessel_list(self):
        """Get list of available vessels from database"""
        return self.db_integration.get_all_vessels()

# Test the enhanced processor
if __name__ == "__main__":
    processor = EnhancedDocumentProcessor()
    
    # Test getting vessel list
    vessels = processor.get_vessel_list()
    print(f"Available vessels: {len(vessels)}")
    
    for vessel in vessels[:5]:  # Show first 5
        print(f"  - {vessel.get('name', 'Unknown')} (IMO: {vessel.get('imo', 'N/A')})")
