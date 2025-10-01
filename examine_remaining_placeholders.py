#!/usr/bin/env python3
"""
Examine Remaining Placeholders
Detailed analysis of what square bracket placeholders remain unfilled
"""

from docx import Document
import re

def examine_placeholders(file_path):
    """Examine remaining placeholders in detail"""
    try:
        doc = Document(file_path)
        
        # Extract all text from the document
        total_text = ""
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            total_text += paragraph.text + "\n"
        
        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_text += paragraph.text + "\n"
        
        # Find all square bracket patterns
        square_patterns = re.findall(r'\[[^\]]*\]', total_text)
        
        # Filter out [Data Not Available] patterns
        unfilled_patterns = [p for p in square_patterns if p != '[Data Not Available]']
        
        print(f"📊 Found {len(unfilled_patterns)} unfilled square bracket patterns:")
        print()
        
        # Group similar patterns
        pattern_counts = {}
        for pattern in unfilled_patterns:
            if pattern in pattern_counts:
                pattern_counts[pattern] += 1
            else:
                pattern_counts[pattern] = 1
        
        # Sort by frequency
        sorted_patterns = sorted(pattern_counts.items(), key=lambda x: x[1], reverse=True)
        
        print("🔍 PATTERN FREQUENCY ANALYSIS:")
        for pattern, count in sorted_patterns:
            print(f"  {count:2d}x: {pattern}")
        
        print()
        
        # Show context for some patterns
        print("📄 CONTEXT EXAMPLES:")
        lines = total_text.split('\n')
        shown_patterns = set()
        
        for i, line in enumerate(lines):
            for pattern in unfilled_patterns[:10]:  # Show context for first 10 unique patterns
                if pattern in line and pattern not in shown_patterns:
                    print(f"  Line {i+1}: ...{line.strip()}...")
                    shown_patterns.add(pattern)
                    if len(shown_patterns) >= 5:  # Limit to 5 examples
                        break
            if len(shown_patterns) >= 5:
                break
        
        return unfilled_patterns
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        return []

def main():
    """Examine the remaining placeholders"""
    print("🔍 EXAMINING REMAINING PLACEHOLDERS")
    print("=" * 50)
    
    test_file = "outputs/test_cleanup_filled.docx"
    
    print(f"📄 Analyzing: {test_file}")
    print()
    
    patterns = examine_placeholders(test_file)
    
    if patterns:
        print()
        print("💡 RECOMMENDATIONS:")
        
        # Analyze pattern types
        tech_specs = [p for p in patterns if any(word in p.lower() for word in ['cetane', 'viscosity', 'ash', 'water', 'sulfur', 'density', 'flash'])]
        contact_info = [p for p in patterns if any(word in p.lower() for word in ['tel', 'phone', 'email', 'fax', 'address'])]
        product_info = [p for p in patterns if any(word in p.lower() for word in ['product', 'quantity', 'price', 'origin', 'destination'])]
        
        if tech_specs:
            print(f"  🔬 Technical specifications: {len(tech_specs)} patterns")
            print(f"     Examples: {tech_specs[:3]}")
        
        if contact_info:
            print(f"  📞 Contact information: {len(contact_info)} patterns")
            print(f"     Examples: {contact_info[:3]}")
        
        if product_info:
            print(f"  📦 Product information: {len(product_info)} patterns")
            print(f"     Examples: {product_info[:3]}")
        
        other = [p for p in patterns if p not in tech_specs and p not in contact_info and p not in product_info]
        if other:
            print(f"  ❓ Other patterns: {len(other)} patterns")
            print(f"     Examples: {other[:3]}")

if __name__ == "__main__":
    main()