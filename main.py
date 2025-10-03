#!/usr/bin/env python3
"""
Simple FastAPI service for document processing
Connects to React platform to process templates with vessel data
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import uvicorn
import uuid
import os
import json
from pathlib import Path
from docx import Document
import re
import shutil
import random
from datetime import datetime, timedelta
from dotenv import load_dotenv
from supabase import create_client, Client
import subprocess

# Load environment variables
load_dotenv()

# Initialize Supabase client
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_ANON_KEY")

if not SUPABASE_URL or not SUPABASE_KEY:
    print("WARNING: Supabase credentials not found. Using dummy data.")
    supabase = None
else:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
    print("SUCCESS: Connected to Supabase database")

app = FastAPI(title="Document Processing Service", version="1.0.0")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Template storage
templates_storage = []
templates_file = Path("templates_data.json")

def load_templates():
    """Load templates from file"""
    global templates_storage
    if templates_file.exists():
        try:
            with open(templates_file, 'r', encoding='utf-8') as f:
                templates_storage = json.load(f)
        except:
            templates_storage = []

def save_templates():
    """Save templates to file"""
    try:
        with open(templates_file, 'w', encoding='utf-8') as f:
            json.dump(templates_storage, f, indent=2)
    except:
        pass

# Load templates on startup
load_templates()

async def get_vessel_data(vessel_imo: str):
    """Get vessel data from database"""
    try:
        if not supabase:
            return get_dummy_vessel_data(vessel_imo)
        
        # Fetch vessel data from database
        response = supabase.table('vessels').select('*').eq('imo', vessel_imo).execute()
        
        if response.data and len(response.data) > 0:
            vessel = response.data[0]
            return map_vessel_data(vessel)
        else:
            return get_dummy_vessel_data(vessel_imo)
            
    except Exception as e:
        print(f"Error fetching vessel data: {e}")
        return get_dummy_vessel_data(vessel_imo)

def map_vessel_data(vessel):
    """Map vessel data to placeholders"""
    return {
        # Vessel data
        'imo_number': vessel.get('imo', 'IMO1234567'),
        'vessel_name': vessel.get('name', 'Sample Vessel'),
        'flag_state': vessel.get('flag', 'Malta'),
        'vessel_type': vessel.get('vessel_type', 'Crude Oil Tanker'),
        'deadweight': vessel.get('deadweight', '75000'),
        'gross_tonnage': vessel.get('gross_tonnage', '45000'),
        'year_built': vessel.get('built', '2015'),
        'length': vessel.get('length', '200.5'),
        'beam': vessel.get('beam', '32.2'),
        'draft': vessel.get('draft', '12.5'),
        'speed': vessel.get('speed', '14.5'),
        
        # Company data
        'buyer_company_name': 'Petroleum Trading Ltd.',
        'buyer_address': '123 Marina Bay, Singapore 018956',
        'buyer_email': 'buyer@petroleumtrading.com',
        'buyer_tel': '+65 6123 4567',
        'seller_company': 'Global Oil Trading Co.',
        'seller_address': '456 Oil Street, Rotterdam, Netherlands',
        'seller_email': 'seller@globaloil.com',
        'seller_tel': '+31 20 123 4567',
        
        # Port data
        'port_loading': 'Port Klang, Malaysia',
        'port_discharge': 'Singapore Port',
        'origin': 'Malaysia',
        'country_of_origin': 'Malaysia',
        
        # Financial data
        'price': 'USD 85.50',
        'total_amount': 'USD 4,275,000.00',
        'payment_terms': 'LC at sight',
        'transaction_currency': 'USD',
        'quantity': '50,000 MT',
        
        # Product data
        'commodity': 'Crude Oil',
        'product_description': 'Malaysian Light Crude Oil',
        'specification': 'API 35-40, Sulfur < 0.5%',
        
        # Date data
        'issue_date': datetime.now().strftime('%Y-%m-%d'),
        'date': datetime.now().strftime('%Y-%m-%d'),
        'validity': (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d'),
        
        # Document data
        'document_number': f"DOC-{datetime.now().year}-{random.randint(1000, 9999)}",
        'invoice_no': f"INV-{datetime.now().year}-{random.randint(1000, 9999)}",
        
        # Bank data
        'buyer_bank_name': 'DBS Bank Ltd.',
        'buyer_swift': 'DBSSSGSG',
        'seller_bank_name': 'ABN AMRO Bank N.V.',
        'seller_swift': 'ABNANL2A',
    }

def get_dummy_vessel_data(vessel_imo):
    """Generate dummy vessel data"""
    return map_vessel_data({
        'imo': vessel_imo,
        'name': 'Sample Vessel',
        'flag': 'Malta',
        'vessel_type': 'Crude Oil Tanker',
        'deadweight': '75000',
        'gross_tonnage': '45000',
        'built': '2015',
        'length': '200.5',
        'beam': '32.2',
        'draft': '12.5',
        'speed': '14.5'
    })

def extract_placeholders_from_docx(file_path):
    """Extract placeholders from Word document"""
    placeholders = set()
    
    try:
        doc = Document(file_path)
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            matches = re.findall(r'\{([^}]+)\}', text)
            placeholders.update(matches)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        matches = re.findall(r'\{([^}]+)\}', text)
                        placeholders.update(matches)
        
        return list(placeholders)
        
    except Exception as e:
        print(f"Error extracting placeholders: {e}")
        return []

def fill_word_template(template_path, output_path, vessel_data):
    """Fill Word template with data"""
    try:
        # Copy template to output
        shutil.copy2(template_path, output_path)
        
        # Open the document
        doc = Document(output_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder in vessel_data:
                if f"{{{placeholder}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{placeholder}}}", str(vessel_data[placeholder]))
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder in vessel_data:
                            if f"{{{placeholder}}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace(f"{{{placeholder}}}", str(vessel_data[placeholder]))
        
        # Save the document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error filling template: {e}")
        return False

def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        # Try LibreOffice first (Linux)
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf', 
            '--outdir', os.path.dirname(pdf_path), docx_path
        ], capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            return True
            
    except Exception as e:
        print(f"LibreOffice conversion failed: {e}")
    
    return False

@app.get("/")
async def root():
    return {"message": "Document Processing Service is running", "status": "ok"}

@app.get("/health")
async def health():
    return {"status": "healthy", "service": "document-processing"}

@app.get("/templates")
async def get_templates():
    """Get list of available templates"""
    return templates_storage

@app.post("/upload-template")
async def upload_template(file: UploadFile = File(...)):
    """Upload a new template"""
    try:
        # Create templates directory
        templates_dir = Path("templates")
        templates_dir.mkdir(exist_ok=True)
        
        # Save uploaded file
        file_path = templates_dir / file.filename
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Extract placeholders
        placeholders = extract_placeholders_from_docx(file_path)
        
        # Create template info
        template_info = {
            "id": str(uuid.uuid4()),
            "name": file.filename.replace('.docx', ''),
            "description": f"Template with {len(placeholders)} placeholders",
            "file_name": file.filename,
            "file_size": len(content),
            "placeholders": placeholders,
            "is_active": True,
            "created_at": datetime.now().isoformat()
        }
        
        # Save to storage
        templates_storage.append(template_info)
        save_templates()
        
        return JSONResponse({
            "success": True,
            "message": "Template uploaded successfully",
            "template": template_info
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Template upload failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.post("/process-document")
async def process_document(
    template_id: str = Form(...),
    vessel_imo: str = Form(...)
):
    """Process document with vessel data"""
    try:
        # Find template
        template = None
        for t in templates_storage:
            if t["id"] == template_id:
                template = t
                break
        
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Get vessel data
        vessel_data = await get_vessel_data(vessel_imo)
        
        # Create outputs directory
        outputs_dir = Path("outputs")
        outputs_dir.mkdir(exist_ok=True)
        
        # Generate document ID
        document_id = str(uuid.uuid4())
        
        # Process template
        template_path = Path("templates") / template["file_name"]
        docx_output = outputs_dir / f"{document_id}.docx"
        pdf_output = outputs_dir / f"{document_id}.pdf"
        
        # Fill template
        if not fill_word_template(template_path, docx_output, vessel_data):
            raise HTTPException(status_code=500, detail="Failed to fill template")
        
        # Convert to PDF
        if not convert_docx_to_pdf(docx_output, pdf_output):
            print("PDF conversion failed, DOCX will be available")
        
        return JSONResponse({
            "success": True,
            "message": "Document processed successfully",
            "document_id": document_id,
            "docx_available": True,
            "pdf_available": pdf_output.exists()
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Document processing failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.get("/download/{document_id}")
async def download_document(document_id: str, format: str = "pdf"):
    """Download processed document"""
    try:
        outputs_dir = Path("outputs")
        
        if format.lower() == "pdf":
            file_path = outputs_dir / f"{document_id}.pdf"
            if not file_path.exists():
                # Try DOCX if PDF not available
                file_path = outputs_dir / f"{document_id}.docx"
                if not file_path.exists():
                    raise HTTPException(status_code=404, detail="Document not found")
        else:
            file_path = outputs_dir / f"{document_id}.docx"
            if not file_path.exists():
                raise HTTPException(status_code=404, detail="Document not found")
        
        return FileResponse(
            path=file_path,
            filename=file_path.name,
            media_type='application/octet-stream'
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    print("Starting Document Processing Service...")
    uvicorn.run(app, host="0.0.0.0", port=port)
