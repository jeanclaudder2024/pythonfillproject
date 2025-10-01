#!/usr/bin/env python3
"""
Batch Processing Script for All Uploaded Word Documents
Processes all Word documents in the uploads folder with realistic random data
"""

import os
import glob
from datetime import datetime
from document_processor import DocumentProcessor
from docx import Document

def get_all_word_documents():
    """Get all Word documents from the uploads folder"""
    uploads_dir = "uploads"
    if not os.path.exists(uploads_dir):
        print(f"❌ Uploads directory '{uploads_dir}' not found!")
        return []
    
    # Find all .docx files
    word_files = glob.glob(os.path.join(uploads_dir, "*.docx"))
    return word_files

def extract_file_info(file_path):
    """Extract file ID and original name from the file path"""
    filename = os.path.basename(file_path)
    
    # Try to extract the UUID part (first 36 characters after removing extension)
    name_without_ext = filename.replace('.docx', '')
    
    # Check if it starts with a UUID pattern
    if len(name_without_ext) > 36 and name_without_ext[8] == '-':
        file_id = name_without_ext[:36]  # Extract UUID
        original_name = name_without_ext[37:]  # Rest is original name
    else:
        # Fallback: use the whole filename as both ID and name
        file_id = name_without_ext
        original_name = name_without_ext
    
    return file_id, original_name

def process_single_document(processor, file_path, file_id, original_name):
    """Process a single document and return results"""
    try:
        print(f"📄 Processing: {original_name}")
        
        # Generate a realistic vessel IMO for this document
        vessel_imo = f"IMO{7000000 + hash(file_id) % 3000000}"  # Generate consistent IMO
        
        # Process the document
        word_output, pdf_output = processor.process_document(file_path, vessel_imo, file_id)
        
        # Count placeholders that were processed
        doc = Document(word_output)
        total_text = ""
        for paragraph in doc.paragraphs:
            total_text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_text += paragraph.text + "\n"
        
        # Check for any remaining unfilled placeholders
        unfilled_curly = total_text.count('{') + total_text.count('}')
        unfilled_square = total_text.count('[') + total_text.count(']') - total_text.count('[Data Not Available]') * 2
        data_not_available = total_text.count('[Data Not Available]')
        
        return {
            'status': 'success',
            'original_name': original_name,
            'file_id': file_id,
            'vessel_imo': vessel_imo,
            'word_output': word_output,
            'pdf_output': pdf_output,
            'unfilled_curly': unfilled_curly,
            'unfilled_square': unfilled_square,
            'data_not_available': data_not_available,
            'file_size_kb': round(os.path.getsize(word_output) / 1024, 2)
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'original_name': original_name,
            'file_id': file_id,
            'error': str(e)
        }

def main():
    """Main batch processing function"""
    print("🚀 Starting Batch Processing of All Uploaded Documents")
    print("=" * 60)
    
    # Initialize the document processor
    print("🔧 Initializing document processor...")
    processor = DocumentProcessor()
    
    # Get all Word documents
    word_files = get_all_word_documents()
    
    if not word_files:
        print("❌ No Word documents found in uploads folder!")
        return
    
    print(f"📋 Found {len(word_files)} Word documents to process")
    print()
    
    # Process each document
    results = []
    successful = 0
    failed = 0
    
    for i, file_path in enumerate(word_files, 1):
        file_id, original_name = extract_file_info(file_path)
        
        print(f"[{i}/{len(word_files)}] Processing: {original_name}")
        print(f"    File ID: {file_id}")
        
        result = process_single_document(processor, file_path, file_id, original_name)
        results.append(result)
        
        if result['status'] == 'success':
            successful += 1
            print(f"    ✅ Success! Generated: {os.path.basename(result['word_output'])}")
            print(f"    📊 File size: {result['file_size_kb']} KB")
            if result['unfilled_curly'] > 0 or result['unfilled_square'] > 0:
                print(f"    ⚠️  Unfilled placeholders detected!")
            if result['data_not_available'] > 0:
                print(f"    ⚠️  {result['data_not_available']} [Data Not Available] found")
        else:
            failed += 1
            print(f"    ❌ Failed: {result['error']}")
        
        print()
    
    # Generate summary report
    print("📊 BATCH PROCESSING SUMMARY")
    print("=" * 60)
    print(f"📄 Total documents processed: {len(word_files)}")
    print(f"✅ Successful: {successful}")
    print(f"❌ Failed: {failed}")
    print(f"📈 Success rate: {(successful/len(word_files)*100):.1f}%")
    print()
    
    if successful > 0:
        print("✅ SUCCESSFUL DOCUMENTS:")
        for result in results:
            if result['status'] == 'success':
                print(f"  📄 {result['original_name']}")
                print(f"      🆔 File ID: {result['file_id']}")
                print(f"      🚢 Vessel IMO: {result['vessel_imo']}")
                print(f"      📁 Word: {os.path.basename(result['word_output'])}")
                print(f"      📁 PDF: {os.path.basename(result['pdf_output'])}")
                print(f"      📊 Size: {result['file_size_kb']} KB")
                
                # Quality indicators
                quality_issues = []
                if result['unfilled_curly'] > 0:
                    quality_issues.append(f"Unfilled curly braces: {result['unfilled_curly']}")
                if result['unfilled_square'] > 0:
                    quality_issues.append(f"Unfilled square brackets: {result['unfilled_square']}")
                if result['data_not_available'] > 0:
                    quality_issues.append(f"Data not available: {result['data_not_available']}")
                
                if quality_issues:
                    print(f"      ⚠️  Issues: {', '.join(quality_issues)}")
                else:
                    print(f"      ✅ Perfect: All placeholders filled!")
                print()
    
    if failed > 0:
        print("❌ FAILED DOCUMENTS:")
        for result in results:
            if result['status'] == 'error':
                print(f"  📄 {result['original_name']}")
                print(f"      🆔 File ID: {result['file_id']}")
                print(f"      ❌ Error: {result['error']}")
                print()
    
    print("🎉 Batch processing completed!")
    print(f"📁 All output files are available in the 'outputs' folder")

if __name__ == "__main__":
    main()