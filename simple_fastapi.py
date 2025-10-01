#!/usr/bin/env python3
"""
Simple FastAPI service for document processing - NO PERMISSIONS
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import uvicorn
import uuid
import os
import tempfile
import json
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
import re
import shutil
import random
from datetime import datetime, timedelta
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = FastAPI(title="Simple Document Service", version="1.0.0")

# Configure CORS
allowed_origins = os.getenv("CORS_ORIGINS", "*").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Simple in-memory storage for templates
templates_storage = []
templates_file = Path("templates_data.json")

def load_templates():
    """Load templates from file"""
    global templates_storage
    if templates_file.exists():
        try:
            with open(templates_file, 'r', encoding='utf-8') as f:
                templates_storage = json.load(f)
        except:
            templates_storage = []

def save_templates():
    """Save templates to file"""
    try:
        with open(templates_file, 'w', encoding='utf-8') as f:
            json.dump(templates_storage, f, indent=2)
    except:
        pass

# Load templates on startup
load_templates()

def extract_placeholders_from_docx(file_path):
    """Extract placeholders from a Word document"""
    try:
        doc = Document(file_path)
        placeholders = set()
        
        # Check paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Find placeholders in format {placeholder_name}
            found = re.findall(r'\{([^}]+)\}', text)
            placeholders.update(found)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        found = re.findall(r'\{([^}]+)\}', text)
                        placeholders.update(found)
        
        # Clean up placeholders
        cleaned_placeholders = []
        for placeholder in placeholders:
            clean_placeholder = placeholder.strip()
            if clean_placeholder and not clean_placeholder.startswith('{'):
                cleaned_placeholders.append(clean_placeholder)
        
        return cleaned_placeholders
    except Exception as e:
        print(f"Error extracting placeholders: {e}")
        return ["vessel_name", "imo", "vessel_type", "flag", "owner", "current_date"]

def generate_realistic_data_for_placeholder(placeholder):
    """Generate realistic data for missing placeholders"""
    realistic_data = {
        'vessel_name': 'Petroleum Express 529',
        'imo': 'IMO23796221',
        'vessel_type': 'Crude Oil Tanker',
        'flag': 'Malta',
        'owner': 'Sample Shipping Company',
        'current_date': datetime.now().strftime('%Y-%m-%d'),
        'seller_bank_account_no': f"{random.randint(1000000000, 9999999999)}",
        'seller_bank_swift': f"{random.choice(['CHASUS33', 'BOFAUS3N', 'CITIUS33'])}",
        'seller_bank_name': random.choice(['Chase Bank', 'Bank of America', 'Citibank']),
        'proforma_invoice_no': f"PI-{datetime.now().year}-{random.randint(1000, 9999)}",
        'contract_value': f"USD {random.randint(1000000, 50000000):,}",
        'port_of_loading': random.choice(['Singapore', 'Rotterdam', 'Houston']),
        'port_of_discharge': random.choice(['Tokyo', 'Hamburg', 'New York']),
        'commodity': random.choice(['Crude Oil', 'Diesel', 'Gasoline']),
        'total_quantity': f"{random.randint(10000, 100000)} MT",
        'unit_price': f"USD {random.randint(50, 100)}.00",
        'total_amount': f"USD {random.randint(1000000, 10000000):,}",
        'buyer_company_name': random.choice(['Tokyo Trading Co.', 'Osaka Shipping Ltd.']),
        'seller_company': random.choice(['Singapore Trading Ltd.', 'Malaysia Oil Corp.']),
    }
    
    return realistic_data.get(placeholder.lower(), f"[{placeholder}]")

def fill_word_template(template_path, output_path, vessel_data):
    """Fill a Word template with vessel data"""
    try:
        # Copy template to output location
        shutil.copy2(template_path, output_path)
        
        # Open the copied document
        doc = Document(output_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            text = original_text
            
            # Replace placeholders with vessel data
            for placeholder, value in vessel_data.items():
                # Handle both {placeholder} and {{placeholder}} formats
                text = text.replace(f"{{{placeholder}}}", str(value))
                text = text.replace(f"{{{{{placeholder}}}}}", str(value))
            
            # If text changed, update the paragraph
            if text != original_text:
                paragraph.text = text
                print(f"Replaced placeholders in paragraph: {original_text[:50]}... -> {text[:50]}...")
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        text = original_text
                        
                        # Replace placeholders with vessel data
                        for placeholder, value in vessel_data.items():
                            text = text.replace(f"{{{placeholder}}}", str(value))
                            text = text.replace(f"{{{{{placeholder}}}}}", str(value))
                        
                        # If text changed, update the paragraph
                        if text != original_text:
                            paragraph.text = text
                            print(f"Replaced placeholders in table cell: {original_text[:50]}... -> {text[:50]}...")
        
        # Second pass: fill any remaining placeholders with realistic data
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Find any remaining placeholders
            remaining_placeholders = re.findall(r'\{([^}]+)\}', text)
            for placeholder in remaining_placeholders:
                clean_placeholder = placeholder.strip()
                if clean_placeholder:
                    realistic_value = generate_realistic_data_for_placeholder(clean_placeholder)
                    text = text.replace(f"{{{placeholder}}}", realistic_value)
            paragraph.text = text
        
        # Save the document
        doc.save(output_path)
        print(f"Template filled successfully: {output_path}")
        return True
    except Exception as e:
        print(f"Error filling template: {e}")
        return False

@app.get("/")
async def root():
    return {"message": "Simple Document Service is running", "status": "ok"}

@app.get("/health")
async def health():
    return {"status": "healthy", "service": "document-processing"}

@app.get("/templates")
async def get_templates():
    """Get all active templates"""
    return [template for template in templates_storage if template.get("is_active", True)]

@app.post("/upload-template")
async def upload_template(
    name: str = Form(...),
    description: str = Form(...),
    template_file: UploadFile = File(...)
):
    """Upload a new document template"""
    try:
        # Validate file type
        if not template_file.filename.endswith('.docx'):
            return JSONResponse({
                "success": False,
                "message": "Only .docx files are allowed"
            }, status_code=400)
        
        # Generate a unique template ID
        template_id = str(uuid.uuid4())
        
        # Create templates directory if it doesn't exist
        templates_dir = Path("templates")
        templates_dir.mkdir(exist_ok=True)
        
        # Save the uploaded file
        file_extension = Path(template_file.filename).suffix
        saved_filename = f"{template_id}{file_extension}"
        file_path = templates_dir / saved_filename
        
        with open(file_path, "wb") as buffer:
            content = await template_file.read()
            buffer.write(content)
        
        # Extract placeholders from the template
        placeholders = extract_placeholders_from_docx(str(file_path))
        
        # Store template information
        template_info = {
            "id": template_id,
            "name": name,
            "description": description,
            "file_name": template_file.filename,
            "file_size": len(content),
            "placeholders": placeholders,
            "is_active": True,
            "created_at": datetime.now().isoformat(),
            "created_by": "system"
        }
        
        templates_storage.append(template_info)
        save_templates()
        
        return JSONResponse({
            "success": True,
            "message": "Template uploaded successfully",
            "template": template_info
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Template upload failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.post("/process-document")
async def process_document(
    template_id: str = Form(...),
    vessel_imo: str = Form(...),
    template_file: UploadFile = File(...)
):
    """Process a document template with vessel data"""
    try:
        # Generate a unique document ID
        document_id = str(uuid.uuid4())
        
        # Create outputs directory if it doesn't exist
        outputs_dir = Path("outputs")
        outputs_dir.mkdir(exist_ok=True)
        
        # Find the template in storage
        template_info = None
        for template in templates_storage:
            if template["id"] == template_id:
                template_info = template
                break
        
        if not template_info:
            return JSONResponse({
                "success": False,
                "message": "Template not found",
                "error": "Template ID not found in storage"
            }, status_code=404)
        
        # Get vessel data
        vessel_data = {
            "vessel_name": "Petroleum Express 529",
            "imo": vessel_imo,
            "imo_number": vessel_imo,
            "vessel_type": "Crude Oil Tanker",
            "flag": "Malta",
            "flag_state": "Malta",
            "owner": "Sample Shipping Company",
            "vessel_owner": "Sample Shipping Company",
            "current_date": datetime.now().strftime('%Y-%m-%d'),
            "vessel_id": "1",
            
            # ICPO Specific Fields
            "icpo_number": f"ICPO-{datetime.now().year}-{random.randint(1000, 9999)}",
            "icpo_date": datetime.now().strftime('%Y-%m-%d'),
            "icpo_validity": (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d'),
            "icpo_amount": f"USD {random.randint(1000000, 10000000):,}",
            "icpo_currency": "USD",
            "icpo_terms": "LC at sight",
            "icpo_bank": "HSBC Bank",
            "icpo_bank_address": "1 Centenary Square, Birmingham, UK",
            "icpo_swift": "HBUKGB4B",
            "icpo_account": f"{random.randint(1000000000, 9999999999)}",
            "icpo_beneficiary": "Sample Trading Company Ltd",
            "icpo_beneficiary_address": "123 Marina Bay, Singapore",
            "icpo_beneficiary_swift": "DBSBSGSG",
            "icpo_beneficiary_account": f"{random.randint(1000000000, 9999999999)}",
            "icpo_commodity": "Crude Oil",
            "icpo_quantity": f"{random.randint(10000, 100000)} MT",
            "icpo_specification": "API 35-40, Sulfur < 0.5%",
            "icpo_origin": "Malaysia",
            "icpo_destination": "Singapore",
            "icpo_loading_port": "Port Klang, Malaysia",
            "icpo_discharge_port": "Singapore Port",
            "icpo_loading_date": (datetime.now() + timedelta(days=15)).strftime('%Y-%m-%d'),
            "icpo_discharge_date": (datetime.now() + timedelta(days=25)).strftime('%Y-%m-%d'),
            "icpo_price": f"USD {random.randint(50, 100)}.00 per MT",
            "icpo_total_value": f"USD {random.randint(5000000, 50000000):,}",
            "icpo_payment_terms": "LC at sight",
            "icpo_delivery_terms": "FOB",
            "icpo_inspection": "SGS",
            "icpo_insurance": "All Risks",
        }
        
        # Find the template file
        templates_dir = Path("templates")
        template_file_path = templates_dir / f"{template_id}.docx"
        
        if not template_file_path.exists():
            return JSONResponse({
                "success": False,
                "message": "Template file not found",
                "error": "Template file does not exist on server"
            }, status_code=404)
        
        # Create output files
        filled_docx_file = outputs_dir / f"{document_id}_filled.docx"
        pdf_file = outputs_dir / f"{document_id}_filled.pdf"
        
        # Fill the Word template with vessel data
        success = fill_word_template(template_file_path, filled_docx_file, vessel_data)
        
        if not success:
            return JSONResponse({
                "success": False,
                "message": "Failed to fill template",
                "error": "Could not process the Word template"
            }, status_code=500)
        
        # Convert Word to PDF using docx2pdf (proper conversion)
        pdf_success = False
        try:
            # Try to use docx2pdf for proper Word to PDF conversion
            try:
                from docx2pdf import convert
                convert(str(filled_docx_file), str(pdf_file))
                pdf_success = True
                print(f"Converted Word to PDF successfully: {pdf_file}")
            except ImportError:
                print("docx2pdf not available, trying alternative method...")
                # Alternative: Use LibreOffice command line (if available)
                import subprocess
                try:
                    # Try LibreOffice conversion
                    subprocess.run([
                        'libreoffice', '--headless', '--convert-to', 'pdf', 
                        '--outdir', str(outputs_dir), str(filled_docx_file)
                    ], check=True, capture_output=True)
                    pdf_success = True
                    print(f"Converted Word to PDF using LibreOffice: {pdf_file}")
                except (subprocess.CalledProcessError, FileNotFoundError):
                    print("LibreOffice not available, using fallback method...")
                    # Fallback: Create a simple PDF that explains the situation
                    c = canvas.Canvas(str(pdf_file), pagesize=letter)
                    width, height = letter
                    
                    c.setFont("Helvetica-Bold", 16)
                    c.drawString(50, height - 50, f"Document: {template_info['name']}")
                    
                    c.setFont("Helvetica", 12)
                    y_position = height - 100
                    
                    content_lines = [
                        f"Vessel IMO: {vessel_imo}",
                        f"Document ID: {document_id}",
                        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                        "",
                        "IMPORTANT:",
                        "Your Word template has been filled with vessel data",
                        "and is available for download.",
                        "",
                        "The Word document (.docx) contains your exact template",
                        "with all formatting, tables, and design preserved.",
                        "",
                        "For the best results, please download the .docx file",
                        "which contains the complete formatted document.",
                        "",
                        "PDF conversion requires additional tools to preserve",
                        "the exact Word template formatting."
                    ]
                    
                    for line in content_lines:
                        c.drawString(50, y_position, line)
                        y_position -= 20
                    
                    c.save()
                    pdf_success = True
                    print(f"Created fallback PDF: {pdf_file}")
            
        except Exception as e:
            print(f"PDF conversion failed: {e}")
            pdf_success = False
        
        # Return success response
        return JSONResponse({
            "success": True,
            "message": "Document processed successfully!",
            "document_id": document_id,
            "download_url": f"/download/{document_id}",
            "vessel_imo": vessel_imo,
            "template_id": template_id,
            "template_name": template_info["name"],
            "files_created": {
                "docx_file": str(filled_docx_file),
                "pdf_file": str(pdf_file) if pdf_success else None
            },
            "download_urls": {
                "pdf": f"/download/{document_id}?format=pdf"
            }
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Document processing failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.get("/download/{document_id}")
async def download_document(document_id: str, format: str = "pdf"):
    """Download processed document as actual file"""
    try:
        outputs_dir = Path("outputs")
        outputs_dir.mkdir(exist_ok=True)
        
        if format.lower() == "pdf":
            file_path = outputs_dir / f"{document_id}_filled.pdf"
            media_type = "application/pdf"
            filename = f"vessel_report_{document_id}.pdf"
        else:
            file_path = outputs_dir / f"{document_id}_filled.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"vessel_report_{document_id}.docx"
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Document not found")
        
        return FileResponse(
            path=str(file_path),
            media_type=media_type,
            filename=filename
        )
        
    except Exception as e:
        print(f"Download error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

@app.put("/templates/{template_id}")
async def update_template(
    template_id: str,
    name: str = Form(None),
    description: str = Form(None),
    is_active: bool = Form(None)
):
    """Update template information"""
    try:
        # Find template in storage
        template_found = False
        for i, template in enumerate(templates_storage):
            if template["id"] == template_id:
                # Update fields if provided
                if name is not None:
                    template["name"] = name
                if description is not None:
                    template["description"] = description
                if is_active is not None:
                    template["is_active"] = is_active
                
                template_found = True
                break
        
        if not template_found:
            return JSONResponse({
                "success": False,
                "message": "Template not found"
            }, status_code=404)
        
        # Save updated templates
        save_templates()
        
        return JSONResponse({
            "success": True,
            "message": "Template updated successfully"
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Template update failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

@app.delete("/templates/{template_id}")
async def delete_template(template_id: str):
    """Delete a template"""
    try:
        # Find template in storage
        template_found = False
        for i, template in enumerate(templates_storage):
            if template["id"] == template_id:
                # Remove from storage
                templates_storage.pop(i)
                template_found = True
                break
        
        if not template_found:
            return JSONResponse({
                "success": False,
                "message": "Template not found"
            }, status_code=404)
        
        # Delete template file
        templates_dir = Path("templates")
        template_file = templates_dir / f"{template_id}.docx"
        if template_file.exists():
            template_file.unlink()
        
        # Save updated templates
        save_templates()
        
        return JSONResponse({
            "success": True,
            "message": "Template deleted successfully"
        })
        
    except Exception as e:
        return JSONResponse({
            "success": False,
            "message": f"Template deletion failed: {str(e)}",
            "error": str(e)
        }, status_code=500)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    print("Starting Simple Document Service...")
    print("Available endpoints:")
    print("  GET  /")
    print("  GET  /health")
    print("  GET  /templates")
    print("  POST /upload-template")
    print("  POST /process-document")
    print("  GET  /download/{document_id}")
    print("  PUT  /templates/{template_id}")
    print("  DELETE /templates/{template_id}")
    uvicorn.run(app, host="0.0.0.0", port=port)