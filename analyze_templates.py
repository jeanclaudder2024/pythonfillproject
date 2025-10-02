#!/usr/bin/env python3
"""
Template Analyzer - Extract placeholders from Word documents
"""

import os
import re
from pathlib import Path
from docx import Document
import json

def extract_placeholders_from_docx(file_path):
    """Extract all placeholders from a Word document"""
    placeholders = set()
    
    try:
        doc = Document(file_path)
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Find placeholders in format {placeholder_name} or {{placeholder_name}}
            matches = re.findall(r'\{([^}]+)\}', text)
            placeholders.update(matches)
            
            # Also check for double braces
            double_matches = re.findall(r'\{\{([^}]+)\}\}', text)
            placeholders.update(double_matches)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        matches = re.findall(r'\{([^}]+)\}', text)
                        placeholders.update(matches)
                        
                        double_matches = re.findall(r'\{\{([^}]+)\}\}', text)
                        placeholders.update(double_matches)
        
        # Clean up placeholders
        cleaned_placeholders = set()
        for placeholder in placeholders:
            # Remove extra spaces and clean up
            cleaned = placeholder.strip()
            if cleaned:
                cleaned_placeholders.add(cleaned)
        
        return list(cleaned_placeholders)
        
    except Exception as e:
        print(f"Error analyzing {file_path}: {e}")
        return []

def analyze_all_templates():
    """Analyze all templates in the template folder"""
    template_folder = Path("template")
    
    if not template_folder.exists():
        print("Template folder not found!")
        return {}
    
    all_placeholders = {}
    
    # Get all .docx files
    docx_files = list(template_folder.glob("*.docx"))
    
    print(f"Found {len(docx_files)} template files:")
    for file in docx_files:
        print(f"  - {file.name}")
    
    print("\n" + "="*60)
    
    for docx_file in docx_files:
        print(f"\nAnalyzing: {docx_file.name}")
        print("-" * 40)
        
        placeholders = extract_placeholders_from_docx(docx_file)
        all_placeholders[docx_file.name] = placeholders
        
        print(f"Found {len(placeholders)} placeholders:")
        for i, placeholder in enumerate(sorted(placeholders), 1):
            print(f"  {i:2d}. {{{placeholder}}}")
    
    return all_placeholders

def categorize_placeholders(all_placeholders):
    """Categorize placeholders by type"""
    categories = {
        'vessel': [],
        'port': [],
        'company': [],
        'refinery': [],
        'date': [],
        'financial': [],
        'product': [],
        'document': [],
        'contact': [],
        'other': []
    }
    
    # Keywords for categorization
    vessel_keywords = ['vessel', 'ship', 'imo', 'flag', 'owner', 'operator', 'captain', 'crew']
    port_keywords = ['port', 'loading', 'discharge', 'departure', 'arrival', 'destination', 'origin']
    company_keywords = ['company', 'buyer', 'seller', 'trader', 'broker', 'agent', 'beneficiary']
    refinery_keywords = ['refinery', 'plant', 'facility', 'terminal']
    date_keywords = ['date', 'time', 'validity', 'expiry', 'eta', 'etd']
    financial_keywords = ['price', 'amount', 'value', 'cost', 'payment', 'currency', 'total', 'fee']
    product_keywords = ['product', 'commodity', 'oil', 'crude', 'fuel', 'cargo', 'specification']
    document_keywords = ['number', 'reference', 'id', 'no', 'code']
    contact_keywords = ['address', 'phone', 'email', 'contact', 'fax', 'tel']
    
    all_unique_placeholders = set()
    for placeholders in all_placeholders.values():
        all_unique_placeholders.update(placeholders)
    
    for placeholder in all_unique_placeholders:
        placeholder_lower = placeholder.lower()
        categorized = False
        
        # Check each category
        for keyword in vessel_keywords:
            if keyword in placeholder_lower:
                categories['vessel'].append(placeholder)
                categorized = True
                break
        
        if not categorized:
            for keyword in port_keywords:
                if keyword in placeholder_lower:
                    categories['port'].append(placeholder)
                    categorized = True
                    break
        
        if not categorized:
            for keyword in company_keywords:
                if keyword in placeholder_lower:
                    categories['company'].append(placeholder)
                    categorized = True
                    break
        
        if not categorized:
            for keyword in refinery_keywords:
                if keyword in placeholder_lower:
                    categories['refinery'].append(placeholder)
                    categorized = True
                    break
        
        if not categorized:
            for keyword in date_keywords:
                if keyword in placeholder_lower:
                    categories['date'].append(placeholder)
                    categorized = True
                    break
        
        if not categorized:
            for keyword in financial_keywords:
                if keyword in placeholder_lower:
                    categories['financial'].append(placeholder)
                    categorized = True
                    break
        
        if not categorized:
            for keyword in product_keywords:
                if keyword in placeholder_lower:
                    categories['product'].append(placeholder)
                    categorized = True
                    break
        
        if not categorized:
            for keyword in document_keywords:
                if keyword in placeholder_lower:
                    categories['document'].append(placeholder)
                    categorized = True
                    break
        
        if not categorized:
            for keyword in contact_keywords:
                if keyword in placeholder_lower:
                    categories['contact'].append(placeholder)
                    categorized = True
                    break
        
        if not categorized:
            categories['other'].append(placeholder)
    
    return categories

def main():
    print("TEMPLATE ANALYZER")
    print("=" * 60)
    
    # Analyze all templates
    all_placeholders = analyze_all_templates()
    
    if not all_placeholders:
        print("No templates found or analyzed!")
        return
    
    print("\n" + "="*60)
    print("CATEGORIZED PLACEHOLDERS")
    print("="*60)
    
    # Categorize placeholders
    categories = categorize_placeholders(all_placeholders)
    
    for category, placeholders in categories.items():
        if placeholders:
            print(f"\n{category.upper()} ({len(placeholders)} placeholders):")
            for placeholder in sorted(placeholders):
                print(f"  - {{{placeholder}}}")
    
    # Save results
    results = {
        'templates': all_placeholders,
        'categories': categories,
        'summary': {
            'total_templates': len(all_placeholders),
            'total_unique_placeholders': len(set().union(*all_placeholders.values())),
            'category_counts': {k: len(v) for k, v in categories.items()}
        }
    }
    
    with open('template_analysis.json', 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    print(f"\nResults saved to: template_analysis.json")
    
    # Print summary
    print("\n" + "="*60)
    print("SUMMARY")
    print("="*60)
    print(f"Total Templates: {results['summary']['total_templates']}")
    print(f"Total Unique Placeholders: {results['summary']['total_unique_placeholders']}")
    print("\nCategory Breakdown:")
    for category, count in results['summary']['category_counts'].items():
        if count > 0:
            print(f"  {category}: {count}")

if __name__ == "__main__":
    main()
