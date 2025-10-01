#!/usr/bin/env python3
"""
Test Replacement Script
Tests the placeholder replacement process to debug why placeholders aren't being replaced
"""

import os
from document_processor import DocumentProcessor
from docx import Document
import uuid

def test_replacement():
    """Test the replacement process on a sample document"""
    print("🧪 TESTING PLACEHOLDER REPLACEMENT")
    print("=" * 60)
    
    # Use one of the uploaded documents
    test_file = "uploads/000d8889-db1e-4707-bc66-a95f78933dc0_Proof_of_Product_Template_with_placeholders.docx"
    
    if not os.path.exists(test_file):
        print(f"❌ Test file not found: {test_file}")
        return
    
    print(f"📄 Testing with: {os.path.basename(test_file)}")
    
    # Initialize processor
    processor = DocumentProcessor()
    
    # Load document
    doc = Document(test_file)
    
    # Find placeholders
    print("\n🔍 FINDING PLACEHOLDERS...")
    placeholders = processor.find_placeholders(doc)
    print(f"Found {len(placeholders)} placeholders:")
    for i, placeholder in enumerate(placeholders[:10]):  # Show first 10
        print(f"  {i+1}. {placeholder}")
    if len(placeholders) > 10:
        print(f"  ... and {len(placeholders) - 10} more")
    
    # Generate replacement data
    print("\n📊 GENERATING REPLACEMENT DATA...")
    replacement_data = processor.generate_random_replacement_data(placeholders)
    print(f"Generated data for {len(replacement_data)} keys:")
    for i, (key, value) in enumerate(list(replacement_data.items())[:10]):  # Show first 10
        print(f"  {i+1}. {key}: {str(value)[:50]}{'...' if len(str(value)) > 50 else ''}")
    if len(replacement_data) > 10:
        print(f"  ... and {len(replacement_data) - 10} more")
    
    # Test replacement on a sample text
    print("\n🔧 TESTING REPLACEMENT ON SAMPLE TEXT...")
    sample_text = "Hello {buyer_company_name}, your order {date} is ready."
    print(f"Original: {sample_text}")
    
    # Manual replacement test
    for placeholder in placeholders:
        if placeholder in replacement_data:
            value = str(replacement_data[placeholder])
            sample_text = sample_text.replace(f"{{{placeholder}}}", value)
    
    print(f"After replacement: {sample_text}")
    
    # Test the actual fill_placeholders_with_random_data method
    print("\n🎯 TESTING ACTUAL FILL METHOD...")
    
    # Create a copy of the document for testing
    test_doc = Document(test_file)
    
    # Get first paragraph text before replacement
    first_paragraph_before = test_doc.paragraphs[0].text if test_doc.paragraphs else "No paragraphs"
    print(f"First paragraph before: {first_paragraph_before[:100]}...")
    
    # Apply replacement
    processor.fill_placeholders_with_random_data(test_doc, placeholders, replacement_data)
    
    # Get first paragraph text after replacement
    first_paragraph_after = test_doc.paragraphs[0].text if test_doc.paragraphs else "No paragraphs"
    print(f"First paragraph after: {first_paragraph_after[:100]}...")
    
    # Check if anything changed
    if first_paragraph_before == first_paragraph_after:
        print("⚠️  NO CHANGES DETECTED - REPLACEMENT FAILED!")
    else:
        print("✅ CHANGES DETECTED - REPLACEMENT WORKING!")
    
    # Save test output
    test_output = f"test_replacement_output_{uuid.uuid4().hex[:8]}.docx"
    test_doc.save(test_output)
    print(f"\n💾 Test output saved as: {test_output}")
    
    # Count remaining placeholders
    print("\n📊 CHECKING REMAINING PLACEHOLDERS...")
    total_text = ""
    for paragraph in test_doc.paragraphs:
        total_text += paragraph.text + "\n"
    for table in test_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total_text += paragraph.text + "\n"
    
    import re
    remaining_curly = re.findall(r'\{[^}]+\}', total_text)
    print(f"Remaining curly placeholders: {len(remaining_curly)}")
    if remaining_curly:
        print(f"Examples: {remaining_curly[:5]}")

if __name__ == "__main__":
    test_replacement()