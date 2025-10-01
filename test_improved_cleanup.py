#!/usr/bin/env python3
"""
Test Improved Cleanup Patterns
Tests the enhanced malformed placeholder handling on a problematic document
"""

import os
from document_processor import DocumentProcessor
from docx import Document
import re

def count_unfilled_placeholders(file_path):
    """Count unfilled placeholders in a document"""
    try:
        doc = Document(file_path)
        
        # Extract all text from the document
        total_text = ""
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            total_text += paragraph.text + "\n"
        
        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_text += paragraph.text + "\n"
        
        # Count different types of unfilled placeholders
        unfilled = {
            'curly_single': len(re.findall(r'\{[^}]*\}', total_text)),
            'curly_double': len(re.findall(r'\{\{[^}]*\}\}', total_text)),
            'square_single': len([match for match in re.findall(r'\[[^\]]*\]', total_text) if match != '[Data Not Available]']),
            'square_double': len(re.findall(r'\[\[[^\]]*\]\]', total_text)),
            'malformed_curly': len(re.findall(r'\{[^}]*$|^[^{]*\}', total_text, re.MULTILINE)),
            'malformed_square': len(re.findall(r'\[[^\]]*$|^[^\[]*\]', total_text, re.MULTILINE)),
            'incomplete_curly': len(re.findall(r'\{[A-Za-z0-9]*(?=\s|\n|$)', total_text)),
            'incomplete_square': len(re.findall(r'\[[A-Za-z0-9]*(?=\s|\n|$)', total_text)),
        }
        
        total_issues = sum(unfilled.values())
        
        return unfilled, total_issues, total_text
        
    except Exception as e:
        return {'error': str(e)}, 0, ""

def main():
    """Test the improved cleanup on a problematic document"""
    print("🧪 TESTING IMPROVED CLEANUP PATTERNS")
    print("=" * 60)
    
    # Test on one of the most problematic documents
    test_file = "uploads/09ac51d2-59bd-4aae-8dba-9d7c9178b429_ICPO_Template_with_Placeholders.docx"
    
    if not os.path.exists(test_file):
        print(f"❌ Test file not found: {test_file}")
        return
    
    print(f"📄 Testing with: {os.path.basename(test_file)}")
    print()
    
    # Initialize processor
    processor = DocumentProcessor()
    
    # Process the document
    print("🔧 Processing document with improved cleanup...")
    try:
        word_output, pdf_output = processor.process_document(
            test_file, 
            "IMO9999999",  # Test IMO
            "test_cleanup"  # Test ID
        )
        
        print(f"✅ Document processed successfully!")
        print(f"📁 Output: {os.path.basename(word_output)}")
        print()
        
        # Analyze the results
        print("📊 ANALYZING RESULTS...")
        unfilled, total_issues, text_content = count_unfilled_placeholders(word_output)
        
        if 'error' in unfilled:
            print(f"❌ Error analyzing results: {unfilled['error']}")
            return
        
        print(f"📈 Total issues found: {total_issues}")
        print()
        
        if total_issues == 0:
            print("🎉 PERFECT! No unfilled placeholders found!")
        else:
            print("📋 BREAKDOWN OF REMAINING ISSUES:")
            for issue_type, count in unfilled.items():
                if count > 0:
                    print(f"  🔸 {issue_type}: {count}")
        
        print()
        
        # Show a sample of the content to verify it looks good
        print("📄 SAMPLE CONTENT (first 500 characters):")
        print("-" * 40)
        print(text_content[:500])
        print("-" * 40)
        
        # Check for specific patterns we were trying to fix
        problematic_patterns = [
            (r'\{[A-Za-z0-9]*(?=\s|\n|$)', "Incomplete curly patterns"),
            (r'\[[A-Za-z0-9]*(?=\s|\n|$)', "Incomplete square patterns"),
            (r'\{\]', "Mismatched {] patterns"),
            (r'\[\}', "Mismatched [} patterns"),
        ]
        
        print("🔍 CHECKING FOR SPECIFIC PROBLEMATIC PATTERNS:")
        for pattern, description in problematic_patterns:
            matches = re.findall(pattern, text_content)
            if matches:
                print(f"  ⚠️  {description}: {len(matches)} found - {matches[:3]}")
            else:
                print(f"  ✅ {description}: None found")
        
        print()
        print("✅ Test completed!")
        
    except Exception as e:
        print(f"❌ Error processing document: {str(e)}")

if __name__ == "__main__":
    main()