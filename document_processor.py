import os
import re
import requests
from docx import Document
from bs4 import BeautifulSoup
import random
from faker import Faker
import subprocess
from datetime import datetime
from dotenv import load_dotenv
import json
from random_data_generator import RandomDataGenerator

# Load environment variables
load_dotenv()

class DocumentProcessor:
    def __init__(self):
        self.fake = Faker()
        self.vessel_data_cache = {}
        self.random_generator = RandomDataGenerator()
        print("✅ Random data generator initialized successfully")
        
    def process_document(self, input_path, vessel_imo, file_id):
        """Process the document by filling placeholders with random data"""
        # Load the document
        doc = Document(input_path)
        
        # Find all placeholders
        placeholders = self.find_placeholders(doc)
        
        # Generate random data for all placeholders
        random_data = self.generate_random_replacement_data(placeholders, vessel_imo)
        
        # Fill placeholders with random data
        self.fill_placeholders_with_random_data(doc, placeholders, random_data)
        
        # Save the filled Word document
        word_output = os.path.join('outputs', f"{file_id}_filled.docx")
        doc.save(word_output)
        
        # Convert to PDF
        pdf_output = os.path.join('outputs', f"{file_id}_filled.pdf")
        self.convert_to_pdf(word_output, pdf_output)
        
        return word_output, pdf_output
    
    def find_placeholders(self, doc):
        """Find all placeholders in the document (format: {placeholder_name} or {{placeholder_name}})"""
        placeholders = set()
        malformed_patterns = set()
        
        def extract_placeholders_from_text(text):
            """Extract placeholders from text, including malformed ones"""
            # Standard curly brace patterns
            matches_double = re.findall(r'\{\{([^}]+)\}\}', text)
            matches_single = re.findall(r'\{([^}]+)\}', text)
            placeholders.update(matches_double)
            placeholders.update(matches_single)
            
            # Square bracket patterns
            matches_square_double = re.findall(r'\[\[([^\]]+)\]\]', text)
            matches_square_single = re.findall(r'\[([^\]]+)\]', text)
            placeholders.update(matches_square_double)
            placeholders.update(matches_square_single)
            
            # Malformed patterns - incomplete opening braces
            # Pattern: "Name: {" or "Company: {" etc.
            malformed_incomplete = re.findall(r'(\w+):\s*\{(?!\w)', text)
            for match in malformed_incomplete:
                placeholder_name = f"{match.lower()}_value"
                placeholders.add(placeholder_name)
                malformed_patterns.add(f"{match}: {{")
            
            # Malformed patterns - incomplete opening square brackets
            # Pattern: "Name: [" or "Company: [" etc.
            malformed_incomplete_square = re.findall(r'(\w+):\s*\[(?!\w)', text)
            for match in malformed_incomplete_square:
                placeholder_name = f"{match.lower()}_value"
                placeholders.add(placeholder_name)
                malformed_patterns.add(f"{match}: [")
            
            # Pattern: "{ incomplete" - opening brace without closing
            malformed_open = re.findall(r'\{([^}]*?)(?:\n|$)', text)
            for match in malformed_open:
                if match.strip() and not re.search(r'\}', match):
                    # This is an incomplete placeholder
                    clean_name = re.sub(r'[^a-zA-Z0-9_]', '_', match.strip().lower())
                    if clean_name:
                        placeholders.add(clean_name)
                        malformed_patterns.add(f"{{{match}")
            
            # Pattern: "[ incomplete" - opening square bracket without closing
            malformed_open_square = re.findall(r'\[([^\]]*?)(?:\n|$)', text)
            for match in malformed_open_square:
                if match.strip() and not re.search(r'\]', match):
                    # This is an incomplete placeholder
                    clean_name = re.sub(r'[^a-zA-Z0-9_]', '_', match.strip().lower())
                    if clean_name:
                        placeholders.add(clean_name)
                        malformed_patterns.add(f"[{match}")
        
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            extract_placeholders_from_text(paragraph.text)
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    extract_placeholders_from_text(cell.text)
        
        # Store malformed patterns for later fixing
        self.malformed_patterns = malformed_patterns
        
        return list(placeholders)
    
    def get_vessel_data(self, vessel_imo):
        """Get vessel data from online sources or return cached data"""
        if vessel_imo in self.vessel_data_cache:
            return self.vessel_data_cache[vessel_imo]
        
        vessel_data = {}
        
        try:
            # Try to search for vessel information online
            search_query = f"vessel IMO {vessel_imo} ship information"
            vessel_info = self.search_vessel_online(search_query)
            vessel_data.update(vessel_info)
        except Exception as e:
            print(f"Error searching online: {e}")
        
        # Cache the data
        self.vessel_data_cache[vessel_imo] = vessel_data
        return vessel_data
    
    def search_vessel_online(self, query):
        """Search for vessel information online"""
        vessel_data = {}
        
        try:
            # Search using a simple web search
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            # Try MarineTraffic or similar vessel tracking sites
            search_url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
            response = requests.get(search_url, headers=headers, timeout=10)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                # Extract basic information from search results
                # This is a simplified approach - in production, you'd use proper APIs
                text_content = soup.get_text().lower()
                
                # Try to extract vessel type
                if 'container' in text_content:
                    vessel_data['vessel_type'] = 'Container Ship'
                elif 'tanker' in text_content:
                    vessel_data['vessel_type'] = 'Tanker'
                elif 'bulk' in text_content:
                    vessel_data['vessel_type'] = 'Bulk Carrier'
                elif 'cargo' in text_content:
                    vessel_data['vessel_type'] = 'Cargo Ship'
                
        except Exception as e:
            print(f"Error in online search: {e}")
        
        return vessel_data
    
    def generate_ai_vessel_data(self, vessel_imo, placeholders):
        """Use OpenAI to generate realistic vessel data"""
        if not self.use_openai:
            return {}
        
        try:
            # Create a prompt for OpenAI
            placeholder_list = ", ".join(placeholders)
            prompt = f"""
            Generate realistic vessel information for a ship with IMO number {vessel_imo}.
            I need data for these placeholders: {placeholder_list}
            
            Please provide realistic maritime data in JSON format. Use these guidelines:
            - vessel_name: Create a realistic ship name
            - vessel_type: Choose from Container Ship, Tanker, Bulk Carrier, Cargo Ship, Ferry, etc.
            - flag: Use common flag states like Panama, Liberia, Marshall Islands, Singapore, Malta
            - captain: Generate a realistic captain name
            - crew: Typical crew size (15-30 for most vessels)
            - company: Create a realistic shipping company name
            - port: Use major international ports
            - length: Realistic vessel length in meters (100-400m)
            - width: Realistic beam width in meters (20-60m)
            - tonnage: Deadweight tonnage (10,000-200,000 DWT)
            - speed: Service speed in knots (10-25 knots)
            - year: Year built (2000-2023)
            - phone: International maritime phone format
            - email: Professional maritime email
            - address: Realistic company address
            - date: Current date
            - time: Current time
            
            Return only valid JSON without any additional text.
            """
            
            response = self.openai_client.chat.completions.create(
                model=self.openai_model,
                messages=[
                    {"role": "system", "content": "You are a maritime data expert. Generate realistic vessel information in JSON format."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=self.max_tokens,
                temperature=0.7
            )
            
            # Parse the response
            ai_response = response.choices[0].message.content.strip()
            
            # Try to extract JSON from the response
            try:
                # Remove any markdown formatting
                if ai_response.startswith('```json'):
                    ai_response = ai_response[7:]
                if ai_response.endswith('```'):
                    ai_response = ai_response[:-3]
                
                vessel_data = json.loads(ai_response)
                print(f"OpenAI generated vessel data: {vessel_data}")
                return vessel_data
                
            except json.JSONDecodeError as e:
                print(f"Error parsing OpenAI JSON response: {e}")
                print(f"Raw response: {ai_response}")
                return {}
                
        except Exception as e:
            print(f"Error calling OpenAI API: {e}")
            return {}
    
    def fill_placeholders(self, doc, placeholders, vessel_data, vessel_imo):
        """Fill placeholders with real or fake data"""
        replacement_data = self.generate_replacement_data(placeholders, vessel_data, vessel_imo)
        
        def replace_in_text(text):
            """Replace placeholders in text, including malformed patterns"""
            original_text = text
            
            # Standard placeholder replacement
            for placeholder, value in replacement_data.items():
                # Handle both single and double curly braces
                single_curly = f"{{{placeholder}}}"
                double_curly = f"{{{{{placeholder}}}}}"
                
                # Handle both single and double square brackets
                single_square = f"[{placeholder}]"
                double_square = f"[[{placeholder}]]"
                
                if single_curly in text:
                    text = text.replace(single_curly, str(value))
                if double_curly in text:
                    text = text.replace(double_curly, str(value))
                if single_square in text:
                    text = text.replace(single_square, str(value))
                if double_square in text:
                    text = text.replace(double_square, str(value))
            
            # Handle malformed patterns
            if hasattr(self, 'malformed_patterns'):
                for malformed_pattern in self.malformed_patterns:
                    if malformed_pattern in text:
                        # Try to find a suitable replacement
                        if ": {" in malformed_pattern:
                            # Pattern like "Name: {" or "Company: {"
                            field_name = malformed_pattern.split(":")[0].strip()
                            placeholder_key = f"{field_name.lower()}_value"
                            if placeholder_key in replacement_data:
                                text = text.replace(malformed_pattern, f"{field_name}: {replacement_data[placeholder_key]}")
                        elif malformed_pattern.startswith("{") and not malformed_pattern.endswith("}"):
                            # Pattern like "{ incomplete"
                            content = malformed_pattern[1:].strip()
                            clean_name = re.sub(r'[^a-zA-Z0-9_]', '_', content.lower())
                            if clean_name in replacement_data:
                                text = text.replace(malformed_pattern, str(replacement_data[clean_name]))
            
            # Additional cleanup for common malformed patterns
            # Fix patterns like "Name: {" -> "Name: [Generated Value]"
            text = re.sub(r'Name:\s*\{(?!\w)', f'Name: {replacement_data.get("name_value", "John Doe")}', text)
            text = re.sub(r'Company:\s*\{(?!\w)', f'Company: {replacement_data.get("company_value", "ABC Corporation")}', text)
            text = re.sub(r'Designation:\s*\{(?!\w)', f'Designation: {replacement_data.get("designation_value", "Manager")}', text)
            
            # Handle complex nested patterns like "{\n{\nOxidation stability"
            text = re.sub(r'\{\s*\n\s*\{\s*\n([^}]+)', r'\1', text)
            
            # Handle complex nested patterns like "[\n[\nOxidation stability"
            text = re.sub(r'\[\s*\n\s*\[\s*\n([^\]]+)', r'\1', text)
            
            # Handle malformed patterns with incomplete content like "{S", "{Re", etc.
            text = re.sub(r'\{[A-Za-z0-9]*(?=\s|\n|$)', '', text)  # Remove incomplete curly patterns
            text = re.sub(r'\[[A-Za-z0-9]*(?=\s|\n|$)', '', text)  # Remove incomplete square patterns
            
            # Handle standalone "{" or "{}" patterns
            text = re.sub(r'\{\s*\}', '', text)  # Remove empty braces
            text = re.sub(r'\{\s*(?=\n|$)', '', text)  # Remove incomplete opening braces at line end
            
            # Handle standalone "[" or "[]" patterns
            text = re.sub(r'\[\s*\]', '', text)  # Remove empty brackets
            text = re.sub(r'\[\s*(?=\n|$)', '', text)  # Remove incomplete opening brackets at line end
            
            # Handle malformed closing patterns like "{]" or "[}"
            text = re.sub(r'\{\]', '', text)  # Remove mismatched {]
            text = re.sub(r'\[\}', '', text)  # Remove mismatched [}
            
            # Handle large blocks of malformed content (like the long blocks we saw)
            # Remove patterns that start with [ and contain multiple lines but no proper closing
            text = re.sub(r'\[\s*\n(?:[^\[\]]*\n)*[^\[\]]*(?=\n|$)', '', text, flags=re.MULTILINE)
            
            # Only clean up truly malformed patterns (incomplete brackets without proper content)
            # Remove patterns like "{ " or "[ " (just opening bracket with whitespace)
            text = re.sub(r':\s*\{\s*(?=\n|$)', ': [Data Not Available]', text)  # Only empty curly patterns
            text = re.sub(r':\s*\[\s*(?=\n|$)', ': [Data Not Available]', text)  # Only empty square patterns
            
            # Clean up remaining standalone empty patterns
            text = re.sub(r'\{\s*(?=\n|$)', '', text)  # Remove standalone empty curly patterns
            text = re.sub(r'\[\s*(?=\n|$)', '', text)  # Remove standalone empty square patterns
            
            # Handle specific technical specification patterns (both complete and malformed)
            text = re.sub(r'Cetane number[^}]*\{[^}]*', f'Cetane number: {replacement_data.get("cetane", "51.7")}', text)
            text = re.sub(r'Cetane index[^}]*\{[^}]*', f'Cetane index: {replacement_data.get("cetane", "51.7")}', text)
            text = re.sub(r'Viscosity @ 40 °C[^}]*\{[^}]*', f'Viscosity @ 40 °C: {replacement_data.get("viscosity", "2.5 cSt")}', text)
            text = re.sub(r'Ash content[^}]*\{[^}]*', f'Ash content: {replacement_data.get("ash_content", "0.005%")}', text)
            text = re.sub(r'Water content[^}]*\{[^}]*', f'Water content: {replacement_data.get("water_content", "0.025%")}', text)
            text = re.sub(r'Sulfur content[^}]*\{[^}]*', f'Sulfur content: {replacement_data.get("sulfur_content", "0.001%")}', text)
            
            # Handle malformed technical specifications (incomplete patterns)
            text = re.sub(r'Cetane number[^:]*$', f'Cetane number: {replacement_data.get("cetane", "51.7")}', text, flags=re.MULTILINE)
            text = re.sub(r'Cetane index[^:]*$', f'Cetane index: {replacement_data.get("cetane", "51.7")}', text, flags=re.MULTILINE)
            text = re.sub(r'Viscosity @ 40 °C[^:]*$', f'Viscosity @ 40 °C: {replacement_data.get("viscosity", "2.5 cSt")}', text, flags=re.MULTILINE)
            text = re.sub(r'Ash content[^:]*$', f'Ash content: {replacement_data.get("ash_content", "0.005%")}', text, flags=re.MULTILINE)
            text = re.sub(r'Water content[^:]*$', f'Water content: {replacement_data.get("water_content", "0.025%")}', text, flags=re.MULTILINE)
            text = re.sub(r'Sulfur content[^:]*$', f'Sulfur content: {replacement_data.get("sulfur_content", "0.001%")}', text, flags=re.MULTILINE)
            
            return text
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            paragraph.text = replace_in_text(paragraph.text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = replace_in_text(paragraph.text)
    
    def generate_replacement_data(self, placeholders, vessel_data, vessel_imo):
        """Generate replacement data for placeholders"""
        replacement_data = {}
        
        # Try to get AI-generated data first
        if self.use_openai:
            ai_data = self.generate_ai_vessel_data(vessel_imo, placeholders)
            if ai_data:
                # Use AI data as the primary source
                vessel_data.update(ai_data)
        
        for placeholder in placeholders:
            placeholder_lower = placeholder.lower().strip()
            
            # Check if we have real data first
            if placeholder_lower in vessel_data:
                replacement_data[placeholder] = vessel_data[placeholder_lower]
                continue
            
            # Generate fake data based on placeholder name
            if 'imo' in placeholder_lower:
                replacement_data[placeholder] = vessel_imo
            elif 'vessel_name' in placeholder_lower or 'ship_name' in placeholder_lower:
                replacement_data[placeholder] = vessel_data.get('vessel_name', self.fake.company() + " " + random.choice(['Star', 'Ocean', 'Wave', 'Sea', 'Marine']))
            elif 'vessel_type' in placeholder_lower or 'ship_type' in placeholder_lower:
                replacement_data[placeholder] = vessel_data.get('vessel_type', random.choice(['Container Ship', 'Tanker', 'Bulk Carrier', 'Cargo Ship', 'Ferry']))
            elif 'flag' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['Panama', 'Liberia', 'Marshall Islands', 'Singapore', 'Malta', 'Bahamas'])
            elif 'port' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['Singapore', 'Shanghai', 'Rotterdam', 'Hamburg', 'Los Angeles', 'Long Beach'])
            elif 'captain' in placeholder_lower or 'master' in placeholder_lower:
                replacement_data[placeholder] = self.fake.name()
            elif 'date' in placeholder_lower:
                replacement_data[placeholder] = datetime.now().strftime('%Y-%m-%d')
            elif 'time' in placeholder_lower:
                replacement_data[placeholder] = datetime.now().strftime('%H:%M')
            elif 'year' in placeholder_lower:
                replacement_data[placeholder] = str(random.randint(2000, 2023))
            elif 'length' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(100, 400)} meters"
            elif 'width' in placeholder_lower or 'beam' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(20, 60)} meters"
            elif 'tonnage' in placeholder_lower or 'dwt' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(10000, 200000)} DWT"
            elif 'speed' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(10, 25)} knots"
            elif 'crew' in placeholder_lower:
                replacement_data[placeholder] = str(random.randint(15, 30))
            elif 'company' in placeholder_lower or 'owner' in placeholder_lower:
                replacement_data[placeholder] = self.fake.company() + " Shipping"
            elif 'phone' in placeholder_lower:
                replacement_data[placeholder] = self.fake.phone_number()
            elif 'email' in placeholder_lower:
                replacement_data[placeholder] = self.fake.email()
            elif 'address' in placeholder_lower:
                replacement_data[placeholder] = self.fake.address().replace('\n', ', ')
            # Specific shipping/cargo placeholders
            elif 'signatory' in placeholder_lower:
                replacement_data[placeholder] = self.fake.name()
            elif 'title' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['Captain', 'Chief Officer', 'Marine Superintendent', 'Port Agent', 'Operations Manager'])
            elif 'reference' in placeholder_lower or 'ref' in placeholder_lower:
                replacement_data[placeholder] = f"REF-{random.randint(100000, 999999)}"
            elif 'document_number' in placeholder_lower or 'doc_number' in placeholder_lower:
                replacement_data[placeholder] = f"DOC-{random.randint(100000, 999999)}"
            elif 'notary_number' in placeholder_lower:
                replacement_data[placeholder] = f"NOT-{random.randint(100000, 999999)}"
            elif 'registration' in placeholder_lower:
                replacement_data[placeholder] = f"REG-{random.randint(100000, 999999)}"
            elif 'tel' in placeholder_lower or 'telephone' in placeholder_lower:
                replacement_data[placeholder] = self.fake.phone_number()
            elif 'commodity' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['Marine Gas Oil', 'Heavy Fuel Oil', 'Diesel', 'Crude Oil', 'Gasoline'])
            elif 'specification' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['ISO 8217:2017', 'ASTM D975', 'EN 590', 'ISO 3675'])
            elif 'origin' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['Singapore', 'UAE', 'Russia', 'Saudi Arabia', 'Kuwait'])
            elif 'shipping_terms' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['FOB', 'CIF', 'CFR', 'DDP', 'EXW'])
            # Technical specifications
            elif 'gravity' in placeholder_lower or 'api' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(15.0, 45.0):.2f}"
            elif 'density' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(0.8, 1.0):.3f} kg/L"
            elif 'viscosity' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(1.5, 6.0):.2f} cSt"
            elif 'flash_point' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(60, 100)}°C"
            elif 'pour_point' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(-30, 10)}°C"
            elif 'cfpp' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(-20, 5)}°C"
            elif 'water_content' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(0.01, 0.05):.3f}%"
            elif 'ash_content' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(0.001, 0.01):.4f}%"
            elif 'sediment' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(0.001, 0.01):.4f}%"
            elif 'lubricity' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(300, 500)} microns"
            elif 'calorific_value' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(40000, 45000)} kJ/kg"
            elif 'octane_number' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(87, 98)}"
            elif 'nitrogen' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(0.1, 0.5):.2f}%"
            elif 'nickel' in placeholder_lower or 'vanadium' in placeholder_lower or 'sodium' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(0.1, 5.0):.2f} ppm"
            elif 'oxygenates' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(0.1, 2.0):.2f}%"
            elif 'dist_' in placeholder_lower:  # Distillation temperatures
                replacement_data[placeholder] = f"{random.randint(150, 350)}°C"
            # Vessel specifications
            elif 'cargo_capacity' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(5000, 50000)} m³"
            elif 'cargo_tanks' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(6, 20)} tanks"
            elif 'draft' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(8.0, 15.0):.1f} meters"
            elif 'engine_type' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['MAN B&W', 'Wärtsilä', 'Caterpillar', 'Mitsubishi'])
            elif 'class_society' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['DNV GL', 'Lloyd\'s Register', 'ABS', 'Bureau Veritas', 'ClassNK'])
            elif 'ism_manager' in placeholder_lower:
                replacement_data[placeholder] = self.fake.company() + " Management"
            elif 'registry_port' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['Monrovia', 'Valletta', 'Singapore', 'Panama City', 'Nassau'])
            elif 'issue_date' in placeholder_lower:
                replacement_data[placeholder] = datetime.now().strftime('%Y-%m-%d')
            # Handle malformed placeholder types
            elif placeholder_lower == 'name_value' or 'name' in placeholder_lower:
                replacement_data[placeholder] = self.fake.name()
            elif placeholder_lower == 'company_value' or ('company' in placeholder_lower and 'value' in placeholder_lower):
                replacement_data[placeholder] = self.fake.company() + " Ltd."
            elif placeholder_lower == 'designation_value' or ('designation' in placeholder_lower and 'value' in placeholder_lower):
                replacement_data[placeholder] = random.choice(['Manager', 'Director', 'Chief Executive', 'Operations Manager', 'General Manager'])
            elif 'principal_buyer_designation' in placeholder_lower:
                replacement_data[placeholder] = random.choice(['Purchasing Manager', 'Chief Procurement Officer', 'Supply Chain Director', 'Commercial Manager'])
            elif 'cetane' in placeholder_lower:
                replacement_data[placeholder] = f"{random.uniform(45.0, 55.0):.1f}"
            elif 'pumping_capacity' in placeholder_lower:
                replacement_data[placeholder] = f"{random.randint(500, 2000)} m³/h"
            else:
                # Default fallback
                replacement_data[placeholder] = f"[{placeholder}]"
        
        return replacement_data
    
    def convert_to_pdf(self, word_path, pdf_path):
        """Convert Word document to PDF"""
        import platform
        
        # Method 1: Try using docx2pdf with proper COM initialization on Windows
        try:
            if platform.system() == "Windows":
                # Initialize COM for Windows
                import pythoncom
                pythoncom.CoInitialize()
                try:
                    from docx2pdf import convert
                    convert(word_path, pdf_path)
                    print(f"PDF conversion successful with docx2pdf: {pdf_path}")
                    return
                finally:
                    pythoncom.CoUninitialize()
            else:
                # Non-Windows systems
                from docx2pdf import convert
                convert(word_path, pdf_path)
                print(f"PDF conversion successful with docx2pdf: {pdf_path}")
                return
        except Exception as e:
            print(f"Error converting to PDF with docx2pdf: {e}")
        
        # Method 2: Try using LibreOffice command line
        try:
            # Try different LibreOffice executable names
            libreoffice_commands = ['soffice', 'libreoffice', 'C:\\Program Files\\LibreOffice\\program\\soffice.exe']
            
            for cmd in libreoffice_commands:
                try:
                    subprocess.run([
                        cmd, '--headless', '--convert-to', 'pdf', 
                        '--outdir', os.path.dirname(pdf_path), word_path
                    ], check=True, timeout=30, capture_output=True)
                    
                    # Rename the output file if needed
                    base_name = os.path.splitext(os.path.basename(word_path))[0]
                    generated_pdf = os.path.join(os.path.dirname(pdf_path), f"{base_name}.pdf")
                    if os.path.exists(generated_pdf) and generated_pdf != pdf_path:
                        os.rename(generated_pdf, pdf_path)
                    
                    if os.path.exists(pdf_path):
                        print(f"PDF conversion successful with LibreOffice: {pdf_path}")
                        return
                        
                except (subprocess.CalledProcessError, FileNotFoundError):
                    continue
                    
        except Exception as e2:
            print(f"Error converting to PDF with LibreOffice: {e2}")
        
        # Method 3: Create fallback text file
        fallback_path = pdf_path.replace('.pdf', '_fallback.txt')
        with open(fallback_path, 'w') as f:
            f.write("PDF conversion failed. Please download the Word document instead.\n\n")
            f.write("To enable PDF conversion, please install one of the following:\n")
            f.write("1. Microsoft Word (recommended for Windows)\n")
            f.write("2. LibreOffice (free alternative)\n\n")
            f.write("LibreOffice can be downloaded from: https://www.libreoffice.org/download/")
        print(f"PDF conversion failed. Fallback text file created: {fallback_path}")
    
    def generate_random_replacement_data(self, placeholders, vessel_imo=None):
        """Generate random replacement data for all placeholders"""
        replacement_data = {}
        
        # Generate comprehensive random data
        comprehensive_data = self.random_generator.generate_comprehensive_data()
        
        # Add the comprehensive data to replacement data
        replacement_data.update(comprehensive_data)
        
        # Generate specific values for each placeholder
        for placeholder in placeholders:
            if placeholder not in replacement_data:
                # Use the random generator to get a value for this placeholder
                value = self.random_generator.get_random_value(placeholder)
                replacement_data[placeholder] = value
        
        # Add some specific vessel-related data if IMO is provided
        if vessel_imo:
            replacement_data['imo_number'] = vessel_imo
            replacement_data['vessel_imo'] = vessel_imo
        
        return replacement_data
    
    def fill_placeholders_with_random_data(self, doc, placeholders, replacement_data):
        """Fill placeholders in the document with random data"""
        def replace_in_text(text, replacement_data):
            """Replace placeholders in text with random data"""
            if not text:
                return text
            
            # Replace placeholders - both curly braces and square brackets
            for placeholder in placeholders:
                if placeholder in replacement_data:
                    value = str(replacement_data[placeholder])
                    # Replace curly brace patterns
                    text = text.replace(f"{{{{{placeholder}}}}}", value)
                    text = text.replace(f"{{{placeholder}}}", value)
                    # Replace square bracket patterns
                    text = text.replace(f"[[{placeholder}]]", value)
                    text = text.replace(f"[{placeholder}]", value)
            
            # Additional cleanup for common malformed patterns
            # Fix patterns like "Name: {" -> "Name: [Generated Value]"
            text = re.sub(r'Name:\s*\{(?!\w)', f'Name: {replacement_data.get("name", "John Doe")}', text)
            text = re.sub(r'Company:\s*\{(?!\w)', f'Company: {replacement_data.get("company_name", "ABC Corporation")}', text)
            text = re.sub(r'Designation:\s*\{(?!\w)', f'Designation: {replacement_data.get("designation", "Manager")}', text)
            
            # Fix patterns like "Name: [" -> "Name: [Generated Value]"
            text = re.sub(r'Name:\s*\[(?!\w)', f'Name: {replacement_data.get("name", "John Doe")}', text)
            text = re.sub(r'Company:\s*\[(?!\w)', f'Company: {replacement_data.get("company_name", "ABC Corporation")}', text)
            text = re.sub(r'Designation:\s*\[(?!\w)', f'Designation: {replacement_data.get("designation", "Manager")}', text)
            
            # Handle complex nested patterns like "{\n{\nOxidation stability"
            text = re.sub(r'\{\s*\n\s*\{\s*\n([^}]+)', r'\1', text)
            
            # Handle standalone "{" or "{}" patterns
            text = re.sub(r'\{\s*\}', '', text)  # Remove empty braces
            text = re.sub(r'\{\s*(?=\n|$)', '', text)  # Remove incomplete opening braces at line end
            
            # Handle specific technical specification patterns - curly braces
            text = re.sub(r'Cetane number[^}]*\{[^}]*', f'Cetane number: {replacement_data.get("cetane_number", "51.7")}', text)
            text = re.sub(r'Cetane index[^}]*\{[^}]*', f'Cetane index: {replacement_data.get("cetane_number", "51.7")}', text)
            text = re.sub(r'Viscosity @ 40 °C[^}]*\{[^}]*', f'Viscosity @ 40 °C: {replacement_data.get("viscosity", "2.5 cSt")}', text)
            text = re.sub(r'Ash content[^}]*\{[^}]*', f'Ash content: {replacement_data.get("ash_content", "0.005%")}', text)
            text = re.sub(r'Water content[^}]*\{[^}]*', f'Water content: {replacement_data.get("water_content", "0.025%")}', text)
            text = re.sub(r'Sulfur content[^}]*\{[^}]*', f'Sulfur content: {replacement_data.get("sulfur_content", "0.001%")}', text)
            
            # Handle specific technical specification patterns - square brackets
            text = re.sub(r'Cetane number[^\]]*\[[^\]]*', f'Cetane number: {replacement_data.get("cetane_number", "51.7")}', text)
            text = re.sub(r'Cetane index[^\]]*\[[^\]]*', f'Cetane index: {replacement_data.get("cetane_number", "51.7")}', text)
            text = re.sub(r'Viscosity @ 40 °C[^\]]*\[[^\]]*', f'Viscosity @ 40 °C: {replacement_data.get("viscosity", "2.5 cSt")}', text)
            text = re.sub(r'Ash content[^\]]*\[[^\]]*', f'Ash content: {replacement_data.get("ash_content", "0.005%")}', text)
            text = re.sub(r'Water content[^\]]*\[[^\]]*', f'Water content: {replacement_data.get("water_content", "0.025%")}', text)
            text = re.sub(r'Sulfur content[^\]]*\[[^\]]*', f'Sulfur content: {replacement_data.get("sulfur_content", "0.001%")}', text)
            
            return text
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                new_text = replace_in_text(paragraph.text, replacement_data)
                if new_text != paragraph.text:
                    # Clear the paragraph and add the new text
                    paragraph.clear()
                    paragraph.add_run(new_text)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            new_text = replace_in_text(paragraph.text, replacement_data)
                            if new_text != paragraph.text:
                                # Clear the paragraph and add the new text
                                paragraph.clear()
                                paragraph.add_run(new_text)